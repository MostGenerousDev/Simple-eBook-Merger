#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════╗
║              Simple eBook Merger v3.0                             ║
║              github.com/GenerousGreivous/Simple-eBook-Merger      ║
╚══════════════════════════════════════════════════════════════╝

Merges multiple eBook files into a single output file.
Supports EPUB, TXT, HTML, PDF, DOCX, RTF, Markdown, FB2, ODT,
and CBZ inputs. Output formats: EPUB, TXT, HTML, PDF, DOCX.

Wrote this because I got tired of juggling 30+ chapter files on
my eReader whenever an author splits their work into individual
files. Just toss them in a folder, run this, done.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
HOW TO RUN
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Just double-click the script or run it from the terminal:

    python ebook_merger.py          (opens the GUI)
    python ebook_merger.py --nogui  (headless / terminal mode)

  In terminal mode you'll need to edit the CONFIG section below
  before running. The GUI handles everything through the window.

PROFILES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Save your merge settings so you don't have to reconfigure
  everything for recurring jobs. The GUI has Save/Load/Delete
  buttons in the Profiles section at the top.

  From the terminal:
    python ebook_merger.py --save-profile "My Series"
    python ebook_merger.py --profile "My Series"
    python ebook_merger.py --list-profiles

  Profiles are stored as JSON in ~/.ebook_merger/profiles/

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
REQUIREMENTS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Python 3.8+  →  https://python.org/downloads
  (check "Add Python to PATH" during install)

  Required (always needed):
    pip install ebooklib beautifulsoup4 lxml

  Optional (install for extra format support):
    pip install pypdf          # PDF input
    pip install python-docx    # DOCX read/write
    pip install odfpy          # ODT input
    pip install fpdf2          # PDF output
    pip install Pillow         # CBZ input + cover images

  tkinter ships with Python on Windows and Mac by default.
  On Linux you might need:  sudo apt install python3-tk

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SUPPORTED INPUT FORMATS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  .epub  — chapters, formatting, images all carried over
  .txt   — turned into a single chapter, paragraphs kept
  .html  — turned into a chapter, formatting kept
  .htm   — same as .html
  .pdf   — text extracted from each page (needs pypdf)
  .docx  — paragraphs extracted, basic formatting kept (needs python-docx)
  .rtf   — RTF tags stripped, plain text wrapped into HTML
  .md    — Markdown converted to HTML
  .fb2   — FictionBook XML parsed into chapters
  .odt   — OpenDocument text extracted (needs odfpy)
  .cbz   — comic book images extracted as pages (needs Pillow)

  Anything else gets skipped automatically.

SUPPORTED OUTPUT FORMATS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  .epub  — EPUB 3 (default)
  .txt   — plain text
  .html  — single HTML file
  .pdf   — PDF document (needs fpdf2)
  .docx  — Word document (needs python-docx)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
NOTES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  - Files get merged in alphabetical order by filename.
    Prefix with numbers to control order (01_, 02_, etc.)
  - Divider pages go between each source file by default,
    so you know where one ends and the next starts.
  - EPUB inputs keep their chapter structure, TXT and HTML
    each become one chapter.
  - Images from EPUB files are preserved.
  - Output is EPUB 3 by default, works with Calibre, Apple
    Books, Kobo, Kindle (via Send to Kindle / conversion),
    Rockbox, etc.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠  LEGAL DISCLAIMER — PLEASE READ, I'M BEGGING YOU
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  BY DOWNLOADING, VIEWING, THINKING ABOUT, DREAMING ABOUT,
  MENTIONING IN PASSING, RUNNING, COMPILING, DECOMPILING,
  REVERSE-ENGINEERING, SCREEN-SHOTTING, PRINTING OUT AND
  FOLDING INTO A PAPER AIRPLANE, OR OTHERWISE INTERACTING
  WITH THIS SOFTWARE IN ANY WAY — INCLUDING BUT NOT LIMITED
  TO STARING AT IT REAL HARD — YOU ACKNOWLEDGE THAT YOU HAVE
  READ, UNDERSTOOD, MEMORIZED, AND TATTOOED THE FOLLOWING
  TERMS ONTO YOUR SOUL. IF YOU DO NOT AGREE, CLOSE THIS
  IMMEDIATELY, DELETE IT FROM YOUR COMPUTER, BURN THE HARD
  DRIVE, AND BURY THE ASHES AT A CROSSROADS AT MIDNIGHT.

  1. PERSONAL USE ONLY — SERIOUSLY
     This script exists for one reason: so you can merge
     ebook files YOU ALREADY OWN into one file for YOUR OWN
     eyeballs. That's it. If you use it for literally
     anything else, that's a you problem, not a me problem.

  2. COPYRIGHT — DON'T BE THAT GUY
     The stuff you merge is almost definitely copyrighted by
     someone who isn't you. You may NOT, under any
     circumstances, in this universe or any parallel one:
       • Share, upload, email, fax, telegraph, send via
         carrier pigeon, or otherwise transmit the output
       • Sell it, trade it, barter it for goats, or
         monetize it in any way whatsoever
       • Post it online anywhere — not even "just for
         friends," not even on a private Discord, not even
         on a USB stick you "accidentally" leave somewhere
       • Use it for AI training, data mining, or feeding
         to any kind of machine learning model
       • Print it out and wallpaper your bathroom with it
         (actually that one might be fine, but I'm still
         not liable)
     Violating copyright law can get you fined into the
     shadow realm and/or imprisoned. Don't test it.

  3. ABSOLUTELY ZERO WARRANTY — NONE — VOID — NADA
     This software is provided "AS IS" and "AS AVAILABLE"
     and "WITH ALL FAULTS" and "GOOD LUCK" and "DON'T LOOK
     AT ME" — without warranty of any kind, express, implied,
     spiritual, hypothetical, or interdimensional, including
     but not limited to the warranties of merchantability,
     fitness for a particular purpose, non-infringement,
     accuracy, reliability, or the warranty that it won't
     set your computer on fire (it won't, probably, but
     legally speaking I make no promises).

     I tested this thing. It works on my machine. If it
     doesn't work on yours: skill issue. (Kidding. Mostly.
     But still not my fault.)

  4. LIMITATION OF LIABILITY — THE BIG ONE
     TO THE MAXIMUM EXTENT PERMITTED BY EVERY LAW THAT HAS
     EVER BEEN WRITTEN, IS CURRENTLY BEING WRITTEN, OR WILL
     BE WRITTEN IN THE FUTURE BY ANY GOVERNING BODY ON
     EARTH, IN SPACE, OR ON ANY CELESTIAL BODY THAT MAY
     EVENTUALLY BE COLONIZED:

     I am not liable for ANYTHING. Not for:
       • Your computer exploding (it won't)
       • Your files getting corrupted (back them up)
       • Your eReader bricking (not how that works)
       • Your cat walking across the keyboard mid-merge
       • Lost profits, lost data, lost time, lost sleep,
         lost friendships, existential dread, or emotional
         distress caused by a poorly formatted chapter break
       • Any lawsuits, fines, cease-and-desists, angry
         emails, or strongly worded letters you receive
       • Acts of God, acts of war, acts of Congress, acts
         of your IT department, or acts of sheer stupidity
       • Anything else that happens, has happened, or could
         theoretically happen in any timeline

     My total liability to you, for everything, forever,
     across all claims, is exactly $0.00 USD. Not a penny
     more. You paid nothing for this. You get nothing back.
     That's how free software works.

  5. INDEMNIFICATION — YOU PROTECT ME, NOT THE OTHER WAY AROUND
     If someone comes after me because of something YOU did
     with this script, YOU agree to:
       • Pay all my legal fees (and my lawyer isn't cheap)
       • Defend me in court (bring snacks)
       • Cover any damages, settlements, fines, or
         judgments — including the ones that haven't been
         invented yet
       • Basically make me whole, as if your reckless
         behavior never happened
     This applies even if I somehow contributed to the
     problem by, say, writing a script that works too well.

  6. YOUR RESPONSIBILITY — ALL OF IT
     You are a grown human being (or a very talented dog,
     I don't judge) and you bear SOLE, COMPLETE, ABSOLUTE,
     TOTAL, UNQUALIFIED, NO-TAKEBACKS responsibility for:
       • Everything you do with this script
       • Everything that happens as a result
       • Making sure you're not breaking any laws —
         federal, state, local, international, maritime,
         intergalactic, or HOA
       • Any consequences, legal or otherwise, including
         but not limited to imprisonment, fines, social
         ostracism, bad karma, or haunting by the ghost
         of a displeased author
     "I didn't read the disclaimer" is not a defense.
     "I didn't know" is not a defense. "My dog did it"
     is not a defense (even if your dog is the talented
     one mentioned above).

  7. SUPPORT THE AUTHORS — THE MORAL PART
     Look, I built this so reading is more convenient,
     not so people can rip off writers. If you're merging
     chapters from a series you love, MAKE SURE you
     actually paid for it. Authors are real people with
     rent to pay. Support them. Buy their stuff. Leave
     nice reviews. Don't be the reason they quit writing.

  8. SEVERABILITY — THE LEGAL SAFETY NET
     If any part of this disclaimer is found to be
     unenforceable by a court of law (or a particularly
     aggressive homeowners association), the rest of it
     still stands. You don't get to throw out the whole
     thing because one clause was too spicy.

  9. GOVERNING LAW
     This disclaimer is governed by the laws of whatever
     jurisdiction the author happens to reside in at the
     time of the dispute — and no, I'm not telling you
     where that is. If there's a legal conflict, it gets
     resolved wherever I feel like filing. You agreed to
     this. Good luck with that.

  10. FINAL CLAUSE — THE "I REALLY MEAN IT" CLAUSE
      By using this software you acknowledge that:
        • You read all of this (yes, ALL of it)
        • You understood it (or at least tried)
        • You agree to every single word
        • You will not now, or at any point in the
          future, try to hold me responsible for
          anything related to this script
        • You accept that this disclaimer may be
          updated at any time and it's on you to
          check for changes
        • If you somehow find a loophole, congrats,
          but the spirit of this disclaimer still
          applies and you know it
"""

import os
import sys
import uuid
import re
import json
import html as html_module
import threading
import xml.etree.ElementTree as ET
import zipfile
import io
from datetime import datetime

# ── dependency checks ──
# only ebooklib + bs4 + lxml are required now — everything
# else is optional and we fall back gracefully

_missing_required = []

try:
    from ebooklib import epub
except ImportError:
    _missing_required.append("ebooklib")

try:
    from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning
    import warnings
    warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)
except ImportError:
    _missing_required.append("beautifulsoup4")

# lxml is needed by bs4 for proper parsing
try:
    import lxml  # noqa: F401
except ImportError:
    _missing_required.append("lxml")

if _missing_required:
    print(
        "\n  Missing required libraries: " + ", ".join(_missing_required) + "\n"
        "  Fix it by running:\n"
        "    pip install ebooklib beautifulsoup4 lxml\n"
    )
    sys.exit(1)

# ── optional deps — each one gets a flag so we can check later ──

HAS_PYPDF = False
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    pass

HAS_DOCX = False
try:
    import docx as python_docx
    HAS_DOCX = True
except ImportError:
    pass

HAS_ODFPY = False
try:
    from odf.opendocument import load as odf_load
    from odf.text import P as OdfP
    from odf import teletype as odf_teletype
    HAS_ODFPY = True
except ImportError:
    pass

HAS_FPDF2 = False
try:
    from fpdf import FPDF
    HAS_FPDF2 = True
except ImportError:
    pass

HAS_PILLOW = False
try:
    from PIL import Image
    HAS_PILLOW = True
except ImportError:
    pass


# ── inlined striprtf — tiny lib, saves a pip install ──
# ported from the striprtf package (MIT license), trimmed down

def _inlined_rtf_to_text(rtf_text):
    """Strip RTF control words and return plain text.
    This is a simplified version of striprtf.rtf_to_text()."""
    # i literally just grabbed the core logic from striprtf so
    # people don't need to install a whole package for ~80 lines
    pattern = re.compile(
        r"\\([a-z]{1,32})(-?\d{1,10})?[ ]?|\\'([0-9a-f]{2})|\\([^a-z])|([{}])|[\r\n]+|(.)",
        re.I,
    )
    destinations = frozenset((
        'aftncn', 'aftnsep', 'aftnsepc', 'annotation', 'atnauthor',
        'atndate', 'atnicn', 'atnid', 'atnparent', 'atnref', 'atntime',
        'atrfend', 'atrfstart', 'author', 'background', 'bkmkend',
        'bkmkstart', 'blipuid', 'buptim', 'category', 'colorschememapping',
        'colortbl', 'comment', 'company', 'creatim', 'datafield',
        'datastore', 'defchp', 'defpap', 'do', 'doccomm', 'docvar',
        'dptxbxtext', 'ebcend', 'ebcstart', 'factoidname', 'falt',
        'fchars', 'ffdeftext', 'ffentrymcr', 'ffexitmcr', 'ffformat',
        'ffhelptext', 'ffl', 'ffname', 'ffstattext', 'field', 'file',
        'filetbl', 'fldinst', 'fldrslt', 'fldtype', 'fname', 'fontemb',
        'fontfile', 'fonttbl', 'footer', 'footerf', 'footerl', 'footerr',
        'footnote', 'formfield', 'ftncn', 'ftnsep', 'ftnsepc', 'g',
        'generator', 'gridtbl', 'header', 'headerf', 'headerl',
        'headerr', 'hl', 'hlfr', 'hlinkbase', 'hlloc', 'hlsrc', 'hsv',
        'htmltag', 'info', 'keycode', 'keywords', 'latentstyles',
        'lchars', 'levelnumbers', 'leveltext', 'lfolevel', 'linkval',
        'list', 'listlevel', 'listname', 'listoverride', 'listoverridetable',
        'listpicture', 'liststylename', 'listtable', 'listtext',
        'lsdlockedexcept', 'macc', 'maccPr', 'mailmerge', 'maln',
        'malnScr', 'manager', 'margPr', 'mbar', 'mbarPr', 'mbaseJc',
        'mbegChr', 'mborderBox', 'mborderBoxPr', 'mbox', 'mboxPr',
        'mchr', 'mcount', 'mctrlPr', 'md', 'mdeg', 'mdegHide', 'mden',
        'mdiff', 'mdPr', 'me', 'mendChr', 'meqArr', 'meqArrPr', 'mf',
        'mfName', 'mfPr', 'mfunc', 'mfuncPr', 'mgroupChr',
        'mgroupChrPr', 'mgrow', 'mhideBot', 'mhideLeft', 'mhideRight',
        'mhideTop', 'mhtmltag', 'mlim', 'mlimloc', 'mlimlow',
        'mlimlowPr', 'mlimupp', 'mlimuppPr', 'mm', 'mmaddfieldname',
        'mmath', 'mmathPict', 'mmathPr', 'mmaxdist', 'mmc', 'mmcJc',
        'mmconnectstr', 'mmconnectstrdata', 'mmcPr', 'mmcs',
        'mmdatasource', 'mmheadersource', 'mmmailsubject', 'mmodso',
        'mmodsofilter', 'mmodsofldmpdata', 'mmodsomappedname',
        'mmodsoname', 'mmodsorecipdata', 'mmodsosort', 'mmodsosrc',
        'mmodsotable', 'mmodsoudl', 'mmodsoudldata', 'mmodsouniquetag',
        'mmPr', 'mmquery', 'mmr', 'mnary', 'mnaryPr', 'mnoBreak',
        'mnum', 'mobjDist', 'moMath', 'moMathPara', 'moMathParaPr',
        'mopEmu', 'mphant', 'mphantPr', 'mplcHide', 'mpos', 'mr',
        'mrad', 'mradPr', 'mrPr', 'msepChr', 'mshow', 'mshp', 'msPre',
        'msPrePr', 'msSub', 'msSubPr', 'msSubSup', 'msSubSupPr',
        'msSup', 'msSupPr', 'mstrikeBLTR', 'mstrikeH', 'mstrikeTLBR',
        'mstrikeV', 'msub', 'msubHide', 'msup', 'msupHide', 'mtransp',
        'mtype', 'mvertJc', 'mvfmf', 'mvfml', 'mvjc', 'mvtof', 'mvtol',
        'mzeroAsc', 'mzeroDesc', 'mzeroWid', 'nesttableprops',
        'nextfile', 'nonesttables', 'objalias', 'objclass', 'objdata',
        'object', 'objname', 'objsect', 'objtime', 'oldcprops',
        'oldpprops', 'oldsprops', 'oldtprops', 'oleclsid', 'operator',
        'panose', 'password', 'passwordhash', 'pgp', 'pgptbl',
        'picprop', 'pict', 'pn', 'pnseclvl', 'pntext', 'pntxta',
        'pntxtb', 'printim', 'private', 'propname', 'protend',
        'protstart', 'protusertbl', 'pxe', 'result', 'revtbl',
        'revtim', 'rsidtbl', 'rxe', 'shp', 'shpgrp', 'shpinst',
        'shppict', 'shprslt', 'shptxt', 'sn', 'sp', 'staticval',
        'stylesheet', 'subject', 'sv', 'svb', 'tc', 'template',
        'themedata', 'title', 'txe', 'ud', 'upr', 'userprops',
        'wgrffmtfilter', 'windowcaption', 'writereservation',
        'writereservhash', 'xe', 'xform', 'xmlattrname', 'xmlattrvalue',
        'xmlclose', 'xmlname', 'xmlnstbl', 'xmlopen',
    ))
    # special chars
    specialchars = {
        'par': '\n', 'sect': '\n\n', 'page': '\n\n',
        'line': '\n', 'tab': '\t',
        'emdash': '\u2014', 'endash': '\u2013',
        'emspace': '\u2003', 'enspace': '\u2002',
        'qmspace': '\u2005', 'bullet': '\u2022',
        'lquote': '\u2018', 'rquote': '\u2019',
        'ldblquote': '\u201c', 'rdblquote': '\u201d',
    }

    stack = []
    ignorable = False
    ucskip = 1
    curskip = 0
    out = []

    for match in pattern.finditer(rtf_text):
        word, arg, hexval, char, brace, tchar = match.groups()
        if brace:
            curskip = 0
            if brace == '{':
                stack.append((ucskip, ignorable))
            elif brace == '}':
                if stack:
                    ucskip, ignorable = stack.pop()
        elif char:
            curskip = 0
            if char == '~':
                if not ignorable:
                    out.append('\u00a0')
            elif char in '{}\\':
                if not ignorable:
                    out.append(char)
            elif char == '*':
                ignorable = True
        elif word:
            curskip = 0
            if word in destinations:
                ignorable = True
            elif ignorable:
                pass
            elif word in specialchars:
                out.append(specialchars[word])
            elif word == 'uc':
                ucskip = int(arg) if arg else 1
            elif word == 'u':
                c = int(arg) if arg else 0
                if c < 0:
                    c += 0x10000
                out.append(chr(c))
                curskip = ucskip
        elif hexval:
            if curskip > 0:
                curskip -= 1
            elif not ignorable:
                out.append(chr(int(hexval, 16)))
        elif tchar:
            if curskip > 0:
                curskip -= 1
            elif not ignorable:
                out.append(tchar)

    return ''.join(out)


# try to use the real striprtf if installed, else use our inlined version
HAS_STRIPRTF = False
try:
    from striprtf.striprtf import rtf_to_text
    HAS_STRIPRTF = True
except ImportError:
    rtf_to_text = _inlined_rtf_to_text


# ── inlined markdown converter — handles the basics so we don't need
# the markdown package for simple stuff ──

def _inlined_markdown_to_html(text):
    """Convert basic markdown to HTML. Handles headers, bold, italic,
    links, images, code blocks, lists, blockquotes, and horizontal rules.
    Not perfect but good enough for most ebook chapters."""
    lines = text.split('\n')
    html_lines = []
    in_code_block = False
    in_list = False
    list_type = None

    for line in lines:
        # fenced code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                html_lines.append('</code></pre>')
                in_code_block = False
            else:
                html_lines.append('<pre><code>')
                in_code_block = True
            continue

        if in_code_block:
            html_lines.append(html_module.escape(line))
            continue

        # close list if we're not in a list item anymore
        stripped = line.strip()
        is_list_item = bool(re.match(r'^[\*\-\+]\s', stripped)) or bool(re.match(r'^\d+\.\s', stripped))
        if in_list and not is_list_item and stripped:
            html_lines.append(f'</{list_type}>')
            in_list = False
            list_type = None

        # blank lines
        if not stripped:
            if in_list:
                pass  # keep list going through blank lines
            else:
                html_lines.append('')
            continue

        # headers
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            content = _inline_markdown(m.group(2))
            html_lines.append(f'<h{level}>{content}</h{level}>')
            continue

        # horizontal rules
        if re.match(r'^[\-\*_]{3,}\s*$', stripped):
            html_lines.append('<hr/>')
            continue

        # blockquotes
        if stripped.startswith('>'):
            content = _inline_markdown(stripped.lstrip('>').strip())
            html_lines.append(f'<blockquote><p>{content}</p></blockquote>')
            continue

        # unordered list
        m = re.match(r'^[\*\-\+]\s+(.*)', stripped)
        if m:
            if not in_list or list_type != 'ul':
                if in_list:
                    html_lines.append(f'</{list_type}>')
                html_lines.append('<ul>')
                in_list = True
                list_type = 'ul'
            html_lines.append(f'<li>{_inline_markdown(m.group(1))}</li>')
            continue

        # ordered list
        m = re.match(r'^\d+\.\s+(.*)', stripped)
        if m:
            if not in_list or list_type != 'ol':
                if in_list:
                    html_lines.append(f'</{list_type}>')
                html_lines.append('<ol>')
                in_list = True
                list_type = 'ol'
            html_lines.append(f'<li>{_inline_markdown(m.group(1))}</li>')
            continue

        # regular paragraph
        html_lines.append(f'<p>{_inline_markdown(stripped)}</p>')

    if in_code_block:
        html_lines.append('</code></pre>')
    if in_list:
        html_lines.append(f'</{list_type}>')

    return '\n'.join(html_lines)


def _inline_markdown(text):
    """Handle inline markdown: bold, italic, code, links, images."""
    # inline code
    text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
    # images (before links so ![...](...) doesn't match [...](...)
    text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'<img src="\2" alt="\1"/>', text)
    # links
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', text)
    # bold+italic
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', text)
    text = re.sub(r'___(.+?)___', r'<b><i>\1</i></b>', text)
    # bold
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__(.+?)__', r'<b>\1</b>', text)
    # italic
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    text = re.sub(r'_(.+?)_', r'<i>\1</i>', text)
    # strikethrough
    text = re.sub(r'~~(.+?)~~', r'<del>\1</del>', text)
    return text


# try to use the real markdown lib if installed, else use our inlined version
HAS_MARKDOWN = False
try:
    import markdown as md_lib
    HAS_MARKDOWN = True
except ImportError:
    pass


def _md_to_html(text):
    """Convert markdown to HTML using the best available method."""
    if HAS_MARKDOWN:
        try:
            return md_lib.markdown(text, extensions=["extra", "codehilite"])
        except Exception:
            # if codehilite isn't available, try without
            try:
                return md_lib.markdown(text, extensions=["extra"])
            except Exception:
                return md_lib.markdown(text)
    return _inlined_markdown_to_html(text)


# ══════════════════════════════════════════════════════════════
# CONFIG — only matters in --nogui (terminal) mode
# ══════════════════════════════════════════════════════════════
# In GUI mode all of this is set through the window instead.

# Folder with the ebook files you want to merge
INPUT_DIR = r"PUT_YOUR_FOLDER_PATH_HERE"

# Where to save the merged file
OUTPUT_FILE = "merged_book.epub"

# Metadata
BOOK_TITLE  = "My Merged eBook"
BOOK_AUTHOR = "Various Authors"

# Stick a divider page between each source file?
ADD_DIVIDERS = True

# Output format: "epub", "txt", "html", "pdf", "docx"
OUTPUT_FORMAT = "epub"

# Sort order: "smart" (detects numbering, chapters, roman numerals, dates)
#             "alpha" (plain alphabetical A-Z)
SORT_ORDER = "smart"

# ══════════════════════════════════════════════════════════════
# END OF CONFIG
# ══════════════════════════════════════════════════════════════


ALLOWED_EXTENSIONS = {
    ".epub", ".txt", ".html", ".htm", ".pdf", ".docx",
    ".rtf", ".md", ".fb2", ".odt", ".cbz",
}

OUTPUT_FORMATS = ["epub", "txt", "html", "pdf", "docx"]

OUTPUT_EXT_MAP = {
    "epub": ".epub",
    "txt": ".txt",
    "html": ".html",
    "pdf": ".pdf",
    "docx": ".docx",
}

# ── profile keys (the fields we save/load) ──
_PROFILE_KEYS = [
    "name", "input_dir", "output_file", "book_title", "book_author",
    "add_dividers", "output_format", "sort_order", "cover_mode",
    "cover_path", "created", "updated",
]


def _app_dir():
    """Returns ~/.ebook_merger, creating it if needed."""
    base = os.path.join(os.path.expanduser("~"), ".ebook_merger")
    os.makedirs(base, exist_ok=True)
    return base


def _profiles_dir():
    """
    Returns the path to the profiles folder, creating it if needed.
    Uses ~/.ebook_merger/profiles/ so profiles survive across script updates.
    """
    base = os.path.join(os.path.expanduser("~"), ".ebook_merger", "profiles")
    os.makedirs(base, exist_ok=True)
    return base


def _profile_path(name):
    """Turns a profile name into a safe filename."""
    safe = re.sub(r'[^\w\s\-]', '', name).strip().replace(' ', '_')
    if not safe:
        safe = "profile"
    return os.path.join(_profiles_dir(), safe + ".json")


def save_profile(name, settings):
    """
    Saves a merge profile to disk.  *settings* should be a dict with
    keys like input_dir, output_file, book_title, etc.
    """
    today = datetime.now().strftime("%Y-%m-%d")
    path = _profile_path(name)

    # if the file already exists, keep the original creation date
    created = today
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as fh:
                old = json.load(fh)
            created = old.get("created", today)
        except (json.JSONDecodeError, OSError):
            pass

    data = {
        "name": name,
        "input_dir": settings.get("input_dir", ""),
        "output_file": settings.get("output_file", ""),
        "book_title": settings.get("book_title", "My Merged eBook"),
        "book_author": settings.get("book_author", "Various Authors"),
        "add_dividers": bool(settings.get("add_dividers", True)),
        "output_format": settings.get("output_format", "epub").lower(),
        "sort_order": settings.get("sort_order", "smart").lower(),
        "cover_mode": settings.get("cover_mode", "none"),
        "cover_path": settings.get("cover_path", ""),
        "created": created,
        "updated": today,
    }

    with open(path, "w", encoding="utf-8") as fh:
        json.dump(data, fh, indent=2)

    return path


def load_profile(name):
    """
    Loads a saved profile by name.  Returns the dict, or None if
    the profile doesn't exist.
    """
    path = _profile_path(name)
    if not os.path.isfile(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as fh:
            return json.load(fh)
    except (json.JSONDecodeError, OSError):
        return None


def delete_profile(name):
    """Deletes a saved profile.  Returns True if it existed."""
    path = _profile_path(name)
    if os.path.isfile(path):
        os.remove(path)
        return True
    return False


def list_profiles():
    """Returns a list of (name, path) tuples for every saved profile."""
    pdir = _profiles_dir()
    results = []
    for fname in sorted(os.listdir(pdir)):
        if not fname.endswith(".json"):
            continue
        fpath = os.path.join(pdir, fname)
        try:
            with open(fpath, "r", encoding="utf-8") as fh:
                data = json.load(fh)
            results.append((data.get("name", fname[:-5]), fpath))
        except (json.JSONDecodeError, OSError):
            continue
    return results


# ── merge history ──

def _history_path():
    return os.path.join(_app_dir(), "history.json")


def _load_history():
    """Load merge history from disk."""
    path = _history_path()
    if not os.path.isfile(path):
        return []
    try:
        with open(path, "r", encoding="utf-8") as fh:
            data = json.load(fh)
        if isinstance(data, list):
            return data
    except (json.JSONDecodeError, OSError):
        pass
    return []


def _save_history(entries):
    """Save merge history to disk, capped at 100 entries."""
    entries = entries[-100:]  # keep only last 100
    path = _history_path()
    try:
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(entries, fh, indent=2)
    except OSError:
        pass


def _add_history_entry(input_files, output_path, output_format, title, author,
                       dividers, sort_order, cover_mode="none", cover_path=""):
    """Record a merge in history."""
    entries = _load_history()
    entries.append({
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "input_files": [f for _, f in input_files] if input_files else [],
        "input_paths": [p for p, _ in input_files] if input_files else [],
        "output_path": os.path.abspath(output_path),
        "output_format": output_format,
        "title": title,
        "author": author,
        "dividers": dividers,
        "sort_order": sort_order,
        "cover_mode": cover_mode,
        "cover_path": cover_path,
    })
    _save_history(entries)


# ──────────────────────────────────────────────────────────────
# Copy protection
# ──────────────────────────────────────────────────────────────

# CSS injected into EPUB/HTML to make text selection painful
_NOCOPY_CSS = """
/* copy protection */
body, p, span, div, h1, h2, h3, h4, h5, h6 {
  -webkit-user-select: none !important;
  -moz-user-select: none !important;
  -ms-user-select: none !important;
  user-select: none !important;
}
@media print { body { display: none !important; } }
"""


def _protect_epub(book):
    """Make the epub harder to casually copy from."""
    # create a CSS stylesheet item with copy-protection rules
    css_item = epub.EpubItem(
        uid="nocopy_css",
        file_name="style/nocopy.css",
        media_type="text/css",
        content=_NOCOPY_CSS.encode("utf-8"),
    )
    book.add_item(css_item)

    # attach the CSS to every xhtml chapter and add invisible spans
    # that break naive copy-paste
    for item in book.get_items():
        if item.get_type() == 9:  # xhtml
            try:
                # link the nocopy stylesheet
                item.add_link(href="style/nocopy.css", rel="stylesheet", type="text/css")
                # inject invisible spans between sentences
                content = item.get_content().decode("utf-8", errors="replace")
                content = content.replace(". ", ".<span style='font-size:0'> </span> ")
                item.set_content(content.encode("utf-8"))
            except Exception:
                pass

    # add an encryption.xml stub — some readers check this and
    # refuse to let you highlight or copy when it's present
    enc_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<encryption xmlns="urn:oasis:names:tc:opendocument:xmlns:container"\n'
        '            xmlns:enc="http://www.w3.org/2001/04/xmlenc#">\n'
        '</encryption>\n'
    )
    enc_item = epub.EpubItem(
        uid="encryption",
        file_name="META-INF/encryption.xml",
        media_type="application/xml",
        content=enc_xml.encode("utf-8"),
    )
    book.add_item(enc_item)


def _protect_text(text):
    """For plain text there's not a lot we can do, but we can
    mix in zero-width joiners between words to break naive
    search/copy and confuse OCR tools a bit."""
    protected = text.replace(" ", " \u200d")
    return protected


def _protect_html(html_str):
    """Add copy-protection CSS and invisible span noise to HTML output."""
    style_block = f"<style>{_NOCOPY_CSS}</style>"
    if "</head>" in html_str:
        html_str = html_str.replace("</head>", f"{style_block}\n</head>", 1)
    elif "<body" in html_str:
        html_str = html_str.replace("<body", f"{style_block}\n<body", 1)
    # add invisible spans to break copy-paste
    html_str = html_str.replace(". ", ".<span style='font-size:0'> </span> ")
    return html_str


def _protect_pdf(pdf):
    """Set PDF metadata to signal restricted permissions."""
    pdf.set_creator("eBook Merger 3.0")


def _protect_docx(doc):
    """Turn on read-only protection on the docx.
    This isn't bulletproof but stops casual editing/copying."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from lxml import etree
    # add document protection element — read-only, no copying
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    body = doc.element.body
    settings = doc.settings.element
    protect = etree.SubElement(settings, f"{{{ns}}}documentProtection")
    protect.set(f"{{{ns}}}edit", "readOnly")
    protect.set(f"{{{ns}}}enforcement", "1")


# ──────────────────────────────────────────────────────────────
# Smart file ordering
# ──────────────────────────────────────────────────────────────

# Roman numeral lookup — covers I through L which is plenty for chapters
_ROMAN_MAP = {
    "I": 1, "II": 2, "III": 3, "IV": 4, "V": 5,
    "VI": 6, "VII": 7, "VIII": 8, "IX": 9, "X": 10,
    "XI": 11, "XII": 12, "XIII": 13, "XIV": 14, "XV": 15,
    "XVI": 16, "XVII": 17, "XVIII": 18, "XIX": 19, "XX": 20,
    "XXI": 21, "XXII": 22, "XXIII": 23, "XXIV": 24, "XXV": 25,
    "XXVI": 26, "XXVII": 27, "XXVIII": 28, "XXIX": 29, "XXX": 30,
    "XXXI": 31, "XXXII": 32, "XXXIII": 33, "XXXIV": 34, "XXXV": 35,
    "XXXVI": 36, "XXXVII": 37, "XXXVIII": 38, "XXXIX": 39, "XL": 40,
    "XLI": 41, "XLII": 42, "XLIII": 43, "XLIV": 44, "XLV": 45,
    "XLVI": 46, "XLVII": 47, "XLVIII": 48, "XLIX": 49, "L": 50,
}

# Valid roman numeral characters (for quick rejection)
_ROMAN_CHARS = set("IVXLCDM")

# Patterns checked in priority order
_PAT_LEADING_NUM = re.compile(r"^(\d+)")
_PAT_CHAPTER_NUM = re.compile(r"(?:chapter|ch|part|section|sect|episode|ep)[\s._-]*(\d+)", re.I)
_PAT_SERIES_NUM = re.compile(r"(?:book|bk|vol|volume|tome)[\s._-]*(\d+)", re.I)
_PAT_ROMAN = re.compile(
    r"(?:chapter|ch|part|section|sect|episode|ep|book|bk|vol|volume|tome)"
    r"[\s._-]*([IVXLCDM]+)(?:[\s._\-\)]|$)", re.I
)
_PAT_DATE_YMD = re.compile(r"(\d{4})[\s._-](\d{1,2})[\s._-](\d{1,2})")
_PAT_DATE_MDY = re.compile(r"(\d{1,2})[\s._-](\d{1,2})[\s._-](\d{4})")

# For natural sort fallback: split text into numeric and non-numeric chunks
_PAT_NATURAL = re.compile(r"(\d+)")


def _roman_to_int(s):
    """Convert a roman numeral string to int, or None if not valid."""
    upper = s.upper()
    if not upper or not all(c in _ROMAN_CHARS for c in upper):
        return None
    return _ROMAN_MAP.get(upper)


def _extract_sort_key(filename):
    """
    Figure out the best sort key for a filename.
    Returns (priority, numeric_key, natural_key, method_name).

    Lower priority = checked first. If a pattern matches we use its
    numeric_key. natural_key is the fallback for tiebreaking.
    """
    stem = os.path.splitext(filename)[0]

    # 1) Leading number prefix: "01_chapter", "2 - intro", "04.story"
    m = _PAT_LEADING_NUM.match(stem)
    if m:
        return (0, int(m.group(1)), [], "numbered prefix")

    # 2) Chapter/Part/Section number: "Chapter 3", "part_04", "ep12"
    m = _PAT_CHAPTER_NUM.search(stem)
    if m:
        return (1, int(m.group(1)), [], "chapter/part number")

    # 3) Series patterns: "Book 1", "Vol 2", "Volume 03"
    m = _PAT_SERIES_NUM.search(stem)
    if m:
        return (2, int(m.group(1)), [], "series/volume number")

    # 4) Roman numerals after keywords: "Chapter_IV", "Part_II"
    m = _PAT_ROMAN.search(stem)
    if m:
        val = _roman_to_int(m.group(1))
        if val is not None:
            return (3, val, [], "roman numeral")

    # 5) Date-based: "2024-01-15_notes" or "01-15-2024_notes"
    m = _PAT_DATE_YMD.search(stem)
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if 1900 <= y <= 2100 and 1 <= mo <= 12 and 1 <= d <= 31:
            return (4, y * 10000 + mo * 100 + d, [], "date (YYYY-MM-DD)")

    m = _PAT_DATE_MDY.search(stem)
    if m:
        mo, d, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if 1900 <= y <= 2100 and 1 <= mo <= 12 and 1 <= d <= 31:
            return (4, y * 10000 + mo * 100 + d, [], "date (MM-DD-YYYY)")

    # 6) Natural sort fallback: split on numbers so "file2" < "file10"
    chunks = _PAT_NATURAL.split(stem.lower())
    natural = []
    for chunk in chunks:
        if chunk.isdigit():
            natural.append((0, int(chunk)))
        else:
            natural.append((1, chunk))
    return (5, 0, natural, "natural sort")


def smart_sort_files(file_list):
    """
    Takes a list of (path, filename) tuples and returns them in smart
    order. Detects numbered prefixes, chapter/part numbers, roman
    numerals, series patterns, dates, and falls back to natural sort.

    Returns (sorted_list, method_summary) where method_summary is a
    short string describing what sorting method(s) were detected.
    """
    if not file_list:
        return [], "nothing to sort"

    # Build sort keys for each file
    keyed = []
    methods_seen = set()
    for path, name in file_list:
        priority, num_key, natural_key, method = _extract_sort_key(name)
        methods_seen.add(method)
        # Sort by: priority bucket, then numeric key, then natural key,
        # then original filename as final tiebreak
        keyed.append((priority, num_key, natural_key, name.lower(), path, name))

    keyed.sort()

    sorted_list = [(item[4], item[5]) for item in keyed]

    # Summarize what we detected
    if len(methods_seen) == 1:
        summary = f"all files sorted by {next(iter(methods_seen))}"
    else:
        # list methods in a sensible order
        ordered = []
        for m in ["numbered prefix", "chapter/part number", "series/volume number",
                   "roman numeral", "date (YYYY-MM-DD)", "date (MM-DD-YYYY)", "natural sort"]:
            if m in methods_seen:
                ordered.append(m)
        summary = "mixed patterns: " + ", ".join(ordered)

    return sorted_list, summary


# ──────────────────────────────────────────────────────────────
# Core merge logic (shared by GUI and terminal mode)
# ──────────────────────────────────────────────────────────────

def find_ebook_files(folder, sort_order="smart"):
    """Grab every supported file from a folder, sorted by name or smart order."""
    hits = []
    skipped = []
    for name in os.listdir(folder):
        full = os.path.join(folder, name)
        if not os.path.isfile(full):
            continue
        ext = os.path.splitext(name)[1].lower()
        if ext in ALLOWED_EXTENSIONS:
            hits.append((full, name))
        else:
            skipped.append(name)

    if sort_order == "smart":
        hits, sort_summary = smart_sort_files(hits)
    else:
        hits.sort(key=lambda x: x[1].lower())
        sort_summary = "alphabetical (A-Z)"

    return hits, skipped, sort_summary


def find_ebook_files_recursive(folder):
    """Scan a folder recursively for all supported ebook files.
    Returns a list of (full_path, relative_path) tuples."""
    results = []
    for dirpath, _dirnames, filenames in os.walk(folder):
        for name in sorted(filenames):
            ext = os.path.splitext(name)[1].lower()
            if ext in ALLOWED_EXTENSIONS:
                full = os.path.join(dirpath, name)
                rel = os.path.relpath(full, folder)
                results.append((full, rel))
    results.sort(key=lambda x: x[1].lower())
    return results


def _uid(prefix="item"):
    """Short random id so filenames inside the epub don't collide."""
    return f"{prefix}_{uuid.uuid4().hex[:8]}"


def _make_divider(source_name, idx):
    """Builds a simple centered title page to go between files."""
    label = os.path.splitext(source_name)[0]
    xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en" lang="en">\n'
        '<head><title>{t}</title></head>\n'
        '<body style="text-align:center; padding-top:40%;">\n'
        '  <h1 style="font-size:1.8em; margin-bottom:0.3em;">{t}</h1>\n'
        '  <hr style="width:50%; margin:1em auto;"/>\n'
        '</body>\n'
        '</html>\n'
    ).format(t=label)
    ch = epub.EpubHtml(title=label, file_name=f"divider_{idx}.xhtml", lang="en")
    ch.set_content(xhtml.encode("utf-8"))
    return ch


# ── input format readers ──

def _pull_epub_chapters(path, name, book, counter):
    """
    Crack open an epub and pull its chapters + images into the
    merged book. Keeps spine order when possible.
    """
    try:
        src = epub.read_epub(path, options={"ignore_ncx": True})
    except Exception as exc:
        return [], counter, f"couldn't read epub: {exc}"

    added = []
    img_map = {}
    css_map = {}
    docs = []

    for item in src.get_items():
        kind = item.get_type()
        if kind == 9:        # xhtml document / chapter
            docs.append(item)
        elif kind == 6:      # image
            new = f"images/{_uid('img')}_{os.path.basename(item.get_name())}"
            img_map[item.get_name()] = new
            img = epub.EpubImage()
            img.file_name = new
            img.media_type = item.media_type
            img.set_content(item.get_content())
            book.add_item(img)
        elif kind == 3:      # css
            new = f"styles/{_uid('css')}_{os.path.basename(item.get_name())}"
            css_map[item.get_name()] = new
            css = epub.EpubItem(
                file_name=new, media_type="text/css",
                content=item.get_content(),
            )
            book.add_item(css)

    # figure out spine order
    spine_ids = []
    try:
        spine_ids = [sid for sid, _ in src.spine]
    except Exception:
        pass

    by_id = {}
    for item in src.get_items():
        if item.get_id():
            by_id[item.get_id()] = item

    if spine_ids:
        ordered, seen = [], set()
        for sid in spine_ids:
            if sid in by_id and by_id[sid].get_type() == 9:
                ordered.append(by_id[sid])
                seen.add(by_id[sid].get_name())
        for d in docs:
            if d.get_name() not in seen:
                ordered.append(d)
        docs = ordered

    for doc in docs:
        counter += 1
        fname = f"chapter_{counter:04d}.xhtml"

        raw = doc.get_content()
        try:
            html = raw.decode("utf-8", errors="replace")
        except AttributeError:
            html = str(raw)

        # rewrite image + css paths so they point to our renamed copies
        for old, new in img_map.items():
            base = os.path.basename(old)
            html = html.replace(old, new).replace(f"../{old}", new)
            if base != old:
                html = html.replace(base, os.path.basename(new))
        for old, new in css_map.items():
            html = html.replace(old, new).replace(f"../{old}", new)

        # try to fish out a chapter title from the html
        title = None
        try:
            soup = BeautifulSoup(html, "lxml")
            tag = soup.find(["h1", "h2", "h3", "title"])
            if tag:
                title = tag.get_text(strip=True)
        except Exception:
            pass
        if not title:
            title = os.path.basename(doc.get_name())
            title = title.replace(".xhtml", "").replace(".html", "")
            title = title.replace("_", " ").replace("-", " ").title()

        ch = epub.EpubHtml(title=title, file_name=fname, lang="en")
        ch.set_content(html.encode("utf-8"))
        book.add_item(ch)
        added.append((ch, title))

    return added, counter, None


def _txt_to_chapter(path, name, counter):
    """Turn a .txt into a single epub chapter."""
    counter += 1
    label = os.path.splitext(name)[0]
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            raw = fh.read()
    except Exception as exc:
        return None, counter, str(exc)

    paras = []
    for ln in raw.split("\n"):
        ln = ln.rstrip()
        if ln:
            ln = ln.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            paras.append(f"    <p>{ln}</p>")
        else:
            paras.append("    <p><br/></p>")

    words = len(raw.split())
    body = "\n".join(paras)
    xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en" lang="en">\n'
        '<head><title>{t}</title></head>\n'
        '<body>\n  <h1>{t}</h1>\n{b}\n</body>\n</html>\n'
    ).format(t=label, b=body)

    ch = epub.EpubHtml(title=label, file_name=f"chapter_{counter:04d}.xhtml", lang="en")
    ch.set_content(xhtml.encode("utf-8"))
    return (ch, label, words), counter, None


def _html_to_chapter(path, name, counter):
    """Turn an .html/.htm file into a single epub chapter."""
    counter += 1
    label = os.path.splitext(name)[0]
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            raw = fh.read()
    except Exception as exc:
        return None, counter, str(exc)

    try:
        soup = BeautifulSoup(raw, "lxml")
        tag = soup.find(["h1", "h2", "h3", "title"])
        if tag:
            label = tag.get_text(strip=True)
        words = len(soup.get_text().split())
    except Exception:
        words = len(raw.split())

    try:
        soup = BeautifulSoup(raw, "lxml")
        body_tag = soup.find("body")
        inner = body_tag.decode_contents() if body_tag else raw
    except Exception:
        inner = raw

    xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en" lang="en">\n'
        '<head><title>{t}</title></head>\n'
        '<body>\n{b}\n</body>\n</html>\n'
    ).format(t=label, b=inner)

    ch = epub.EpubHtml(title=label, file_name=f"chapter_{counter:04d}.xhtml", lang="en")
    ch.set_content(xhtml.encode("utf-8"))
    return (ch, label, words), counter, None


def _read_pdf(path, name, counter):
    """Extract text from a PDF, one chapter per page (or all as one)."""
    if not HAS_PYPDF:
        return None, counter, "PDF input needs pypdf — pip install pypdf"

    label = os.path.splitext(name)[0]

    try:
        reader = PdfReader(path)
    except Exception as exc:
        return None, counter, f"couldn't read PDF: {exc}"

    all_text = []
    for page_num, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        all_text.append(text)

    combined = "\n\n".join(all_text)
    if not combined.strip():
        return None, counter, "no text found in PDF"

    counter += 1
    words = len(combined.split())

    paras = []
    for ln in combined.split("\n"):
        ln = ln.rstrip()
        if ln:
            ln = html_module.escape(ln)
            paras.append(f"    <p>{ln}</p>")
        else:
            paras.append("    <p><br/></p>")

    body = "\n".join(paras)
    xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en" lang="en">\n'
        '<head><title>{t}</title></head>\n'
        '<body>\n  <h1>{t}</h1>\n{b}\n</body>\n</html>\n'
    ).format(t=label, b=body)

    ch = epub.EpubHtml(title=label, file_name=f"chapter_{counter:04d}.xhtml", lang="en")
    ch.set_content(xhtml.encode("utf-8"))
    return (ch, label, words), counter, None


def _read_docx(path, name, counter):
    """Extract paragraphs from a DOCX file."""
    if not HAS_DOCX:
        return None, counter, "DOCX input needs python-docx — pip install python-docx"

    label = os.path.splitext(name)[0]

    try:
        doc = python_docx.Document(path)
    except Exception as exc:
        return None, counter, f"couldn't read DOCX: {exc}"

    counter += 1
    paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            paras.append("    <p><br/></p>")
            continue

        # basic formatting: bold, italic, headings
        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name:
            paras.append(f"    <h1>{html_module.escape(text)}</h1>")
        elif "heading 2" in style_name:
            paras.append(f"    <h2>{html_module.escape(text)}</h2>")
        elif "heading 3" in style_name:
            paras.append(f"    <h3>{html_module.escape(text)}</h3>")
        else:
            # build inline formatting from runs
            parts = []
            for run in para.runs:
                t = html_module.escape(run.text)
                if run.bold and run.italic:
                    t = f"<b><i>{t}</i></b>"
                elif run.bold:
                    t = f"<b>{t}</b>"
                elif run.italic:
                    t = f"<i>{t}</i>"
                parts.append(t)
            paras.append(f"    <p>{''.join(parts)}</p>")

    full_text = " ".join(p.text for p in doc.paragraphs)
    words = len(full_text.split())

    body = "\n".join(paras)
    xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en" lang="en">\n'
        '<head><title>{t}</title></head>\n'
        '<body>\n  <h1>{t}</h1>\n{b}\n</body>\n</html>\n'
    ).format(t=label, b=body)

    ch = epub.EpubHtml(title=label, file_name=f"chapter_{counter:04d}.xhtml", lang="en")
    ch.set_content(xhtml.encode("utf-8"))
    return (ch, label, words), counter, None


def _read_rtf(path, name, counter):
    """Strip RTF formatting and turn into a chapter."""
    label = os.path.splitext(name)[0]

    try:
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            raw = fh.read()
        text = rtf_to_text(raw)
    except Exception as exc:
        return None, counter, f"couldn't read RTF: {exc}"

    if not text.strip():
        return None, counter, "empty RTF file"

    counter += 1
    words = len(text.split())

    paras = []
    for ln in text.split("\n"):
        ln = ln.rstrip()
        if ln:
            paras.append(f"    <p>{html_module.escape(ln)}</p>")
        else:
            paras.append("    <p><br/></p>")

    body = "\n".join(paras)
    xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en" lang="en">\n'
        '<head><title>{t}</title></head>\n'
        '<body>\n  <h1>{t}</h1>\n{b}\n</body>\n</html>\n'
    ).format(t=label, b=body)

    ch = epub.EpubHtml(title=label, file_name=f"chapter_{counter:04d}.xhtml", lang="en")
    ch.set_content(xhtml.encode("utf-8"))
    return (ch, label, words), counter, None


def _read_markdown(path, name, counter):
    """Convert Markdown to HTML and make a chapter."""
    label = os.path.splitext(name)[0]

    try:
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            raw = fh.read()
        html_content = _md_to_html(raw)
    except Exception as exc:
        return None, counter, f"couldn't read Markdown: {exc}"

    if not raw.strip():
        return None, counter, "empty Markdown file"

    counter += 1
    words = len(raw.split())

    xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en" lang="en">\n'
        '<head><title>{t}</title></head>\n'
        '<body>\n  <h1>{t}</h1>\n{b}\n</body>\n</html>\n'
    ).format(t=label, b=html_content)

    ch = epub.EpubHtml(title=label, file_name=f"chapter_{counter:04d}.xhtml", lang="en")
    ch.set_content(xhtml.encode("utf-8"))
    return (ch, label, words), counter, None


def _read_fb2(path, name, counter):
    """Parse FictionBook2 XML into a chapter."""
    label = os.path.splitext(name)[0]

    try:
        tree = ET.parse(path)
        root = tree.getroot()
    except Exception as exc:
        return None, counter, f"couldn't read FB2: {exc}"

    # FB2 namespace
    ns = ""
    if root.tag.startswith("{"):
        ns = root.tag.split("}")[0] + "}"

    # extract text from all <p> elements in <body>
    body = root.find(f".//{ns}body")
    if body is None:
        return None, counter, "no body found in FB2"

    paras = []
    for p in body.iter(f"{ns}p"):
        text = "".join(p.itertext()).strip()
        if text:
            paras.append(f"    <p>{html_module.escape(text)}</p>")

    if not paras:
        return None, counter, "no text found in FB2"

    # try to get title from <title-info><book-title>
    try:
        ti = root.find(f".//{ns}title-info/{ns}book-title")
        if ti is not None and ti.text:
            label = ti.text.strip()
    except Exception:
        pass

    counter += 1
    full_text = " ".join("".join(p.itertext()) for p in body.iter(f"{ns}p"))
    words = len(full_text.split())

    body_html = "\n".join(paras)
    xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en" lang="en">\n'
        '<head><title>{t}</title></head>\n'
        '<body>\n  <h1>{t}</h1>\n{b}\n</body>\n</html>\n'
    ).format(t=html_module.escape(label), b=body_html)

    ch = epub.EpubHtml(title=label, file_name=f"chapter_{counter:04d}.xhtml", lang="en")
    ch.set_content(xhtml.encode("utf-8"))
    return (ch, label, words), counter, None


def _read_odt(path, name, counter):
    """Extract text from an ODT (OpenDocument Text) file."""
    if not HAS_ODFPY:
        return None, counter, "ODT input needs odfpy — pip install odfpy"

    label = os.path.splitext(name)[0]

    try:
        doc = odf_load(path)
    except Exception as exc:
        return None, counter, f"couldn't read ODT: {exc}"

    paras = []
    for p in doc.getElementsByType(OdfP):
        text = odf_teletype.extractText(p).strip()
        if text:
            paras.append(f"    <p>{html_module.escape(text)}</p>")
        else:
            paras.append("    <p><br/></p>")

    if not any(t.strip() for t in paras):
        return None, counter, "no text found in ODT"

    counter += 1
    full_text = " ".join(
        odf_teletype.extractText(p) for p in doc.getElementsByType(OdfP)
    )
    words = len(full_text.split())

    body = "\n".join(paras)
    xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en" lang="en">\n'
        '<head><title>{t}</title></head>\n'
        '<body>\n  <h1>{t}</h1>\n{b}\n</body>\n</html>\n'
    ).format(t=label, b=body)

    ch = epub.EpubHtml(title=label, file_name=f"chapter_{counter:04d}.xhtml", lang="en")
    ch.set_content(xhtml.encode("utf-8"))
    return (ch, label, words), counter, None


def _read_cbz(path, name, book, counter):
    """
    Extract images from a CBZ (comic book zip) and add each
    as a page in the output.
    """
    if not HAS_PILLOW:
        return [], counter, "CBZ input needs Pillow — pip install Pillow"

    label = os.path.splitext(name)[0]
    added = []

    try:
        with zipfile.ZipFile(path, "r") as zf:
            image_names = sorted([
                n for n in zf.namelist()
                if n.lower().endswith((".jpg", ".jpeg", ".png", ".gif", ".webp"))
                and not n.startswith("__MACOSX")
            ])

            if not image_names:
                return [], counter, "no images found in CBZ"

            for img_name in image_names:
                counter += 1
                img_data = zf.read(img_name)

                # figure out the media type
                ext = os.path.splitext(img_name)[1].lower()
                media_types = {
                    ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                    ".png": "image/png", ".gif": "image/gif",
                    ".webp": "image/webp",
                }
                media_type = media_types.get(ext, "image/jpeg")

                # add image to epub
                img_fname = f"images/{_uid('cbz')}_{os.path.basename(img_name)}"
                img_item = epub.EpubImage()
                img_item.file_name = img_fname
                img_item.media_type = media_type
                img_item.set_content(img_data)
                book.add_item(img_item)

                # create an xhtml page for this image
                page_title = f"{label} - Page {counter}"
                xhtml = (
                    '<?xml version="1.0" encoding="UTF-8"?>\n'
                    '<!DOCTYPE html>\n'
                    '<html xmlns="http://www.w3.org/1999/xhtml"'
                    ' xml:lang="en" lang="en">\n'
                    '<head><title>{t}</title></head>\n'
                    '<body style="text-align:center; margin:0; padding:0;">\n'
                    '  <img src="{src}" alt="{t}"'
                    ' style="max-width:100%; max-height:100vh;"/>\n'
                    '</body>\n</html>\n'
                ).format(t=page_title, src=img_fname)

                ch = epub.EpubHtml(
                    title=page_title,
                    file_name=f"chapter_{counter:04d}.xhtml",
                    lang="en",
                )
                ch.set_content(xhtml.encode("utf-8"))
                book.add_item(ch)
                added.append((ch, page_title))

    except zipfile.BadZipFile:
        return [], counter, "not a valid CBZ/zip file"
    except Exception as exc:
        return [], counter, f"couldn't read CBZ: {exc}"

    return added, counter, None


# ── chapter preview helper ──

def _preview_file_text(path, max_chars=500):
    """Extract the first ~max_chars of text from a file for preview.
    Returns a string of plain text, or an error message."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".epub":
            src = epub.read_epub(path, options={"ignore_ncx": True})
            texts = []
            for item in src.get_items():
                if item.get_type() == 9:
                    raw = item.get_content().decode("utf-8", errors="replace")
                    soup = BeautifulSoup(raw, "lxml")
                    texts.append(soup.get_text(separator="\n"))
                    if len("\n".join(texts)) > max_chars:
                        break
            return "\n".join(texts)[:max_chars]

        elif ext in (".txt", ".md", ".rtf"):
            with open(path, "r", encoding="utf-8", errors="replace") as fh:
                raw = fh.read(max_chars + 200)
            if ext == ".rtf":
                raw = rtf_to_text(raw)
            return raw[:max_chars]

        elif ext in (".html", ".htm"):
            with open(path, "r", encoding="utf-8", errors="replace") as fh:
                raw = fh.read(max_chars * 3)
            soup = BeautifulSoup(raw, "lxml")
            return soup.get_text(separator="\n")[:max_chars]

        elif ext == ".pdf":
            if not HAS_PYPDF:
                return "(PDF preview needs pypdf)"
            reader = PdfReader(path)
            texts = []
            for page in reader.pages:
                texts.append(page.extract_text() or "")
                if len("\n".join(texts)) > max_chars:
                    break
            return "\n".join(texts)[:max_chars]

        elif ext == ".docx":
            if not HAS_DOCX:
                return "(DOCX preview needs python-docx)"
            doc = python_docx.Document(path)
            texts = []
            for para in doc.paragraphs:
                texts.append(para.text)
                if len("\n".join(texts)) > max_chars:
                    break
            return "\n".join(texts)[:max_chars]

        elif ext == ".fb2":
            tree = ET.parse(path)
            root = tree.getroot()
            ns = ""
            if root.tag.startswith("{"):
                ns = root.tag.split("}")[0] + "}"
            body = root.find(f".//{ns}body")
            if body is None:
                return "(no body in FB2)"
            texts = []
            for p in body.iter(f"{ns}p"):
                texts.append("".join(p.itertext()))
                if len("\n".join(texts)) > max_chars:
                    break
            return "\n".join(texts)[:max_chars]

        elif ext == ".odt":
            if not HAS_ODFPY:
                return "(ODT preview needs odfpy)"
            doc = odf_load(path)
            texts = []
            for p in doc.getElementsByType(OdfP):
                texts.append(odf_teletype.extractText(p))
                if len("\n".join(texts)) > max_chars:
                    break
            return "\n".join(texts)[:max_chars]

        elif ext == ".cbz":
            return "(CBZ contains images — no text preview)"

        else:
            return f"(no preview for {ext} files)"

    except Exception as exc:
        return f"(preview error: {exc})"


# ── cover image helpers ──

def _extract_epub_cover(epub_path):
    """Try to pull the cover image bytes from an epub. Returns (data, ext) or (None, None)."""
    try:
        src = epub.read_epub(epub_path, options={"ignore_ncx": True})
        # check metadata for cover image id
        cover_id = None
        for meta in src.get_metadata("OPF", "meta") or []:
            # meta is a tuple like ({attribs}, content)
            if isinstance(meta, tuple) and len(meta) >= 1:
                attrs = meta[0] if isinstance(meta[0], dict) else {}
                if not isinstance(attrs, dict):
                    # sometimes it's (content, attrs)
                    if len(meta) >= 2 and isinstance(meta[1], dict):
                        attrs = meta[1]
                if attrs.get("name") == "cover":
                    cover_id = attrs.get("content")
                    break

        for item in src.get_items():
            if item.get_type() == 6:  # image
                if cover_id and item.get_id() == cover_id:
                    ext = os.path.splitext(item.get_name())[1].lower() or ".jpg"
                    return item.get_content(), ext
        # fallback: grab the first image
        for item in src.get_items():
            if item.get_type() == 6:
                ext = os.path.splitext(item.get_name())[1].lower() or ".jpg"
                return item.get_content(), ext
    except Exception:
        pass
    return None, None


def _extract_cbz_first_image(cbz_path):
    """Grab the first image from a CBZ. Returns (data, ext) or (None, None)."""
    try:
        with zipfile.ZipFile(cbz_path, "r") as zf:
            names = sorted([
                n for n in zf.namelist()
                if n.lower().endswith((".jpg", ".jpeg", ".png", ".gif", ".webp"))
                and not n.startswith("__MACOSX")
            ])
            if names:
                ext = os.path.splitext(names[0])[1].lower()
                return zf.read(names[0]), ext
    except Exception:
        pass
    return None, None


def _get_cover_image_data(cover_mode, cover_path, input_files):
    """Resolve cover image based on mode. Returns (image_bytes, ext) or (None, None).
    cover_mode is 'browse', 'auto', or 'none'."""
    if cover_mode == "none":
        return None, None

    if cover_mode == "browse" and cover_path and os.path.isfile(cover_path):
        try:
            with open(cover_path, "rb") as f:
                data = f.read()
            ext = os.path.splitext(cover_path)[1].lower()
            return data, ext
        except Exception:
            return None, None

    if cover_mode == "auto" and input_files:
        # try the first file — if it's an epub, grab its cover
        first_path = input_files[0][0]
        first_ext = os.path.splitext(first_path)[1].lower()
        if first_ext == ".epub":
            return _extract_epub_cover(first_path)
        elif first_ext == ".cbz":
            return _extract_cbz_first_image(first_path)

    return None, None


# ── output format writers ──

def _strip_html_tags(html_str):
    """Remove HTML tags and return plain text."""
    try:
        soup = BeautifulSoup(html_str, "lxml")
        return soup.get_text(separator="\n")
    except Exception:
        return re.sub(r"<[^>]+>", "", html_str)


def _chapters_to_content_list(chapters_data):
    """
    Normalize chapter data into a list of dicts with
    'title' and 'content' (HTML string) keys.
    """
    result = []
    for item in chapters_data:
        if isinstance(item, dict):
            result.append(item)
        elif isinstance(item, tuple) and len(item) == 2:
            ch, title = item
            try:
                content = ch.get_content().decode("utf-8", errors="replace")
            except Exception:
                content = str(ch.get_content())
            result.append({"title": title, "content": content})
    return result


def _write_epub(chapters_data, spine_items, toc, output_path, title, author,
                dividers_list, book, cover_data=None, cover_ext=None, log=print):
    """Write the merged book as EPUB (the original format)."""
    # add cover image if we have one
    if cover_data and cover_ext:
        media_map = {
            ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".png": "image/png", ".gif": "image/gif",
            ".webp": "image/webp", ".bmp": "image/bmp",
        }
        media_type = media_map.get(cover_ext, "image/jpeg")
        cover_fname = f"images/cover{cover_ext}"

        # add the image item
        cover_img = epub.EpubImage()
        cover_img.file_name = cover_fname
        cover_img.media_type = media_type
        cover_img.set_content(cover_data)
        book.add_item(cover_img)

        # set cover metadata
        book.set_cover(cover_fname, cover_data)

    book.toc = toc
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav"] + spine_items

    _protect_epub(book)

    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    epub.write_epub(output_path, book, {})


def _write_txt(chapters_data, output_path, title, author, log=print):
    """Write merged content as plain text."""
    lines = []
    lines.append(f"{'=' * 60}")
    lines.append(f"  {title}")
    lines.append(f"  by {author}")
    lines.append(f"{'=' * 60}")
    lines.append("")

    content_list = _chapters_to_content_list(chapters_data)
    for i, ch in enumerate(content_list):
        if i > 0:
            lines.append("")
            lines.append(f"{'─' * 40}")
            lines.append("")
        lines.append(f"  {ch['title']}")
        lines.append(f"{'─' * 40}")
        lines.append("")
        text = _strip_html_tags(ch["content"])
        lines.append(text)

    text_out = "\n".join(lines)
    text_out = _protect_text(text_out)

    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text_out)


def _write_html(chapters_data, output_path, title, author,
                cover_data=None, cover_ext=None, log=print):
    """Write merged content as a single HTML file."""
    content_list = _chapters_to_content_list(chapters_data)

    body_parts = []

    # insert cover image as base64 if we have one
    if cover_data and cover_ext:
        import base64
        media_map = {
            ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".png": "image/png", ".gif": "image/gif",
            ".webp": "image/webp", ".bmp": "image/bmp",
        }
        media_type = media_map.get(cover_ext, "image/jpeg")
        b64 = base64.b64encode(cover_data).decode("ascii")
        body_parts.append(
            f'    <div style="text-align:center; margin-bottom:2em;">'
            f'<img src="data:{media_type};base64,{b64}" '
            f'alt="Cover" style="max-width:100%; max-height:80vh;"/></div>'
        )

    for i, ch in enumerate(content_list):
        if i > 0:
            body_parts.append('    <hr style="margin: 2em 0;"/>')
        body_parts.append(f'    <h2>{html_module.escape(ch["title"])}</h2>')

        # extract just the body content from the xhtml
        try:
            soup = BeautifulSoup(ch["content"], "lxml")
            body_tag = soup.find("body")
            inner = body_tag.decode_contents() if body_tag else ch["content"]
        except Exception:
            inner = ch["content"]
        body_parts.append(f"    {inner}")

    html_out = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{html_module.escape(title)}</title>
<style>
  body {{ font-family: Georgia, 'Times New Roman', serif; max-width: 800px;
         margin: 2em auto; padding: 0 1em; line-height: 1.6; color: #222; }}
  h1 {{ text-align: center; border-bottom: 2px solid #333; padding-bottom: 0.5em; }}
  h2 {{ margin-top: 2em; color: #444; }}
  p {{ text-indent: 1.5em; margin: 0.5em 0; }}
  hr {{ border: none; border-top: 1px solid #ccc; }}
</style>
</head>
<body>
  <h1>{html_module.escape(title)}</h1>
  <p style="text-align:center; color:#666;">by {html_module.escape(author)}</p>
{chr(10).join(body_parts)}
</body>
</html>"""

    html_out = _protect_html(html_out)

    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_out)


def _write_pdf(chapters_data, output_path, title, author,
               cover_data=None, cover_ext=None, log=print):
    """Write merged content as a PDF."""
    if not HAS_FPDF2:
        raise RuntimeError("PDF output needs fpdf2 — pip install fpdf2")

    content_list = _chapters_to_content_list(chapters_data)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_title(title)
    pdf.set_author(author)

    _protect_pdf(pdf)

    def _clean_for_pdf(s):
        """Strip zero-width chars and force latin-1 safe text."""
        for zw in "\u200b\u200c\u200d\ufeff":
            s = s.replace(zw, "")
        try:
            s.encode("latin-1")
        except (UnicodeEncodeError, UnicodeDecodeError):
            s = s.encode("ascii", errors="replace").decode("ascii")
        return s

    # cover page if we have an image
    if cover_data and cover_ext and HAS_PILLOW:
        try:
            img = Image.open(io.BytesIO(cover_data))
            # save to a temp file because fpdf wants a path
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=cover_ext, delete=False) as tmp:
                tmp.write(cover_data)
                tmp_path = tmp.name
            pdf.add_page()
            # fit the image to the page
            page_w = pdf.w - pdf.l_margin - pdf.r_margin
            page_h = pdf.h - 30
            img_w, img_h = img.size
            ratio = min(page_w / img_w, page_h / img_h)
            w = img_w * ratio
            h = img_h * ratio
            x = (pdf.w - w) / 2
            pdf.image(tmp_path, x=x, y=15, w=w, h=h)
            os.unlink(tmp_path)
        except Exception:
            pass  # skip cover if it fails

    # title page
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 24)
    pdf.cell(0, 60, _clean_for_pdf(title), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 14)
    pdf.cell(0, 10, _clean_for_pdf(f"by {author}"), new_x="LMARGIN", new_y="NEXT", align="C")

    for ch in content_list:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_x(pdf.l_margin)
        pdf.cell(0, 10, _clean_for_pdf(ch["title"]), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)

        text = _strip_html_tags(ch["content"])
        text = _clean_for_pdf(text)
        pdf.set_font("Helvetica", "", 11)

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(3)
                continue
            # reset x to left margin before each line just in case
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(w=0, h=5, text=line, new_x="LMARGIN", new_y="NEXT")

    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    pdf.output(output_path)


def _write_docx(chapters_data, output_path, title, author,
                cover_data=None, cover_ext=None, log=print):
    """Write merged content as a DOCX file."""
    if not HAS_DOCX:
        raise RuntimeError("DOCX output needs python-docx — pip install python-docx")

    content_list = _chapters_to_content_list(chapters_data)

    doc = python_docx.Document()

    # insert cover image if we have one
    if cover_data and cover_ext:
        try:
            from docx.shared import Inches
            img_stream = io.BytesIO(cover_data)
            doc.add_picture(img_stream, width=Inches(5))
            doc.add_page_break()
        except Exception:
            pass  # skip cover if it fails

    # title
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"by {author}")
    doc.add_paragraph("")

    for i, ch in enumerate(content_list):
        if i > 0:
            doc.add_page_break()
        doc.add_heading(ch["title"], level=1)

        text = _strip_html_tags(ch["content"])
        for line in text.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)

    _protect_docx(doc)

    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    doc.save(output_path)


# ── main merge function ──

def merge_files(files, output_path, title, author, dividers=True,
                output_format="epub", cover_mode="none", cover_path="",
                log=print):
    """
    The actual merge. Takes a list of (path, filename) tuples,
    builds the output in the requested format, writes it to output_path.

    `log` is a callable — in terminal mode it's just print(),
    in GUI mode it pushes text to the log widget.

    Returns a dict with stats or raises on fatal errors.
    """
    # check if the requested output format is available
    fmt = output_format.lower().strip()
    if fmt == "pdf" and not HAS_FPDF2:
        raise RuntimeError("PDF output needs fpdf2 — pip install fpdf2")
    if fmt == "docx" and not HAS_DOCX:
        raise RuntimeError("DOCX output needs python-docx — pip install python-docx")

    # resolve cover image
    cover_data, cover_ext = _get_cover_image_data(cover_mode, cover_path, files)

    # we always build chapters using the epub structures internally,
    # then convert to the target format at the end
    book = epub.EpubBook()
    book.set_identifier(f"merged-{uuid.uuid4().hex[:12]}")
    book.set_title(title)
    book.set_language("en")
    book.add_author(author)
    book.add_metadata("DC", "date", datetime.today().strftime("%Y-%m-%d"))

    spine_items = []
    toc = []
    ch_count = 0
    ok_count = 0
    total_words = 0
    all_chapters = []  # list of (ch_object, title) for output writers

    for i, (fpath, fname) in enumerate(files):
        ext = os.path.splitext(fname)[1].lower()
        log(f"  [{i+1}/{len(files)}] {fname}")

        # check if this format is available
        if ext == ".pdf" and not HAS_PYPDF:
            log(f"    \u26a0 PDF input needs pypdf — pip install pypdf")
            continue
        if ext == ".docx" and not HAS_DOCX:
            log(f"    \u26a0 DOCX input needs python-docx — pip install python-docx")
            continue
        if ext == ".odt" and not HAS_ODFPY:
            log(f"    \u26a0 ODT input needs odfpy — pip install odfpy")
            continue
        if ext == ".cbz" and not HAS_PILLOW:
            log(f"    \u26a0 CBZ input needs Pillow — pip install Pillow")
            continue

        if dividers and i > 0:
            d = _make_divider(fname, i)
            book.add_item(d)
            spine_items.append(d)

        if ext == ".epub":
            chapters, ch_count, err = _pull_epub_chapters(fpath, fname, book, ch_count)
            if err:
                log(f"    \u26a0 {err}")
                continue
            if not chapters:
                log(f"    \u26a0 no chapters found")
                continue

            for ch, _ in chapters:
                spine_items.append(ch)
            all_chapters.extend(chapters)

            if len(chapters) == 1:
                toc.append(chapters[0])
            else:
                sec = os.path.splitext(fname)[0]
                toc.append((epub.Section(sec), [ch for ch, _ in chapters]))

            wc = sum(
                len(BeautifulSoup(ch.get_content(), "lxml").get_text().split())
                for ch, _ in chapters
            )
            total_words += wc
            log(f"    \u2713 {len(chapters)} chapter(s), ~{wc:,} words")
            ok_count += 1

        elif ext == ".cbz":
            chapters, ch_count, err = _read_cbz(fpath, fname, book, ch_count)
            if err:
                log(f"    \u26a0 {err}")
                continue
            if not chapters:
                log(f"    \u26a0 no images found")
                continue

            for ch, _ in chapters:
                spine_items.append(ch)
            all_chapters.extend(chapters)

            if len(chapters) == 1:
                toc.append(chapters[0])
            else:
                sec = os.path.splitext(fname)[0]
                toc.append((epub.Section(sec), [ch for ch, _ in chapters]))

            log(f"    \u2713 {len(chapters)} page(s)")
            ok_count += 1

        else:
            # all single-chapter formats
            reader_map = {
                ".txt":  _txt_to_chapter,
                ".html": _html_to_chapter,
                ".htm":  _html_to_chapter,
                ".pdf":  _read_pdf,
                ".docx": _read_docx,
                ".rtf":  _read_rtf,
                ".md":   _read_markdown,
                ".fb2":  _read_fb2,
                ".odt":  _read_odt,
            }
            reader = reader_map.get(ext)
            if not reader:
                log(f"    \u26a0 unsupported format: {ext}")
                continue

            result, ch_count, err = reader(fpath, fname, ch_count)
            if err or not result:
                log(f"    \u26a0 {err or 'empty file'}")
                continue

            ch, label, wc = result
            book.add_item(ch)
            spine_items.append(ch)
            toc.append(ch)
            all_chapters.append((ch, label))
            total_words += wc
            log(f"    \u2713 1 chapter, ~{wc:,} words")
            ok_count += 1

    if not spine_items:
        raise RuntimeError("Nothing to merge \u2014 no chapters were produced.")

    # write output in the requested format
    if fmt == "epub":
        _write_epub(all_chapters, spine_items, toc, output_path, title, author,
                     [], book, cover_data, cover_ext, log)
    elif fmt == "txt":
        _write_txt(all_chapters, output_path, title, author, log)
    elif fmt == "html":
        _write_html(all_chapters, output_path, title, author,
                     cover_data, cover_ext, log)
    elif fmt == "pdf":
        _write_pdf(all_chapters, output_path, title, author,
                    cover_data, cover_ext, log)
    elif fmt == "docx":
        _write_docx(all_chapters, output_path, title, author,
                     cover_data, cover_ext, log)
    else:
        raise RuntimeError(f"Unknown output format: {fmt}")

    size = os.path.getsize(output_path)
    if size > 1024 * 1024:
        size_nice = f"{size / (1024*1024):.1f} MB"
    else:
        size_nice = f"{size / 1024:.1f} KB"

    return {
        "files_merged": ok_count,
        "chapters": ch_count,
        "words": total_words,
        "size": size_nice,
        "path": os.path.abspath(output_path),
    }


# ──────────────────────────────────────────────────────────────
# Format availability helpers — used by the GUI to show what's
# available based on installed deps
# ──────────────────────────────────────────────────────────────

def _available_input_extensions():
    """Return the set of input extensions that are actually usable."""
    exts = {".epub", ".txt", ".html", ".htm", ".rtf", ".md", ".fb2"}
    if HAS_PYPDF:
        exts.add(".pdf")
    if HAS_DOCX:
        exts.add(".docx")
    if HAS_ODFPY:
        exts.add(".odt")
    if HAS_PILLOW:
        exts.add(".cbz")
    return exts


def _available_output_formats():
    """Return list of output format strings that are actually usable."""
    fmts = ["epub", "txt", "html"]
    if HAS_FPDF2:
        fmts.append("pdf")
    if HAS_DOCX:
        fmts.append("docx")
    return fmts


# ──────────────────────────────────────────────────────────────
# GUI (tkinter)
# ──────────────────────────────────────────────────────────────

def launch_gui():
    """
    Opens the graphical interface. Everything's in here so tkinter
    is only imported when we actually need it — keeps the terminal
    mode dependency-free.
    """
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox
    import json as _json

    # ── theme definitions ──
    # each theme is a dict of colour slots that _apply_theme() maps onto
    # every widget in the window
    THEMES = {
        "Dark": {
            "bg": "#1e1e2e", "bg_mid": "#282840", "bg_entry": "#313148",
            "fg": "#e0e0e8", "fg_dim": "#8888a0", "accent": "#7c6ff7",
            "accent2": "#9b8afb", "green": "#50c878", "red": "#ff6b6b",
            "button_bg": "#3a3a55", "button_fg": "#e0e0e8",
            "heading_fg": "#e0e0e8", "border": "#3a3a55",
            "log_bg": "#282840", "log_fg": "#e0e0e8",
            "select_bg": "#7c6ff7", "warn": "#e8a838",
        },
        "Light": {
            "bg": "#f0f0f5", "bg_mid": "#ffffff", "bg_entry": "#ffffff",
            "fg": "#1a1a2e", "fg_dim": "#6b6b80", "accent": "#5b52d4",
            "accent2": "#7268e0", "green": "#2d8f4e", "red": "#d44040",
            "button_bg": "#dddde8", "button_fg": "#1a1a2e",
            "heading_fg": "#1a1a2e", "border": "#c0c0d0",
            "log_bg": "#ffffff", "log_fg": "#1a1a2e",
            "select_bg": "#5b52d4", "warn": "#c07820",
        },
        "Midnight": {
            "bg": "#0d1b2a", "bg_mid": "#1b2838", "bg_entry": "#1b2838",
            "fg": "#c8dce8", "fg_dim": "#607888", "accent": "#3a86c8",
            "accent2": "#5a9cd8", "green": "#40b868", "red": "#e05555",
            "button_bg": "#1b2838", "button_fg": "#c8dce8",
            "heading_fg": "#d0e4f0", "border": "#2a3a4a",
            "log_bg": "#1b2838", "log_fg": "#c8dce8",
            "select_bg": "#3a86c8", "warn": "#d0a030",
        },
        "Forest": {
            "bg": "#1a2e1a", "bg_mid": "#243524", "bg_entry": "#2a3f2a",
            "fg": "#d0e8d0", "fg_dim": "#78a078", "accent": "#48a848",
            "accent2": "#60c060", "green": "#50d870", "red": "#e06050",
            "button_bg": "#2a3f2a", "button_fg": "#d0e8d0",
            "heading_fg": "#e0f0e0", "border": "#3a5a3a",
            "log_bg": "#243524", "log_fg": "#d0e8d0",
            "select_bg": "#48a848", "warn": "#d0b040",
        },
        "Sunset": {
            "bg": "#2a1a1a", "bg_mid": "#3a2222", "bg_entry": "#452a2a",
            "fg": "#f0dcc8", "fg_dim": "#a08070", "accent": "#d07030",
            "accent2": "#e08840", "green": "#60b860", "red": "#e05050",
            "button_bg": "#452a2a", "button_fg": "#f0dcc8",
            "heading_fg": "#f8e8d0", "border": "#5a3a30",
            "log_bg": "#3a2222", "log_fg": "#f0dcc8",
            "select_bg": "#d07030", "warn": "#e0a030",
        },
    }

    # ── theme persistence ──
    _theme_dir = os.path.join(os.path.expanduser("~"), ".ebook_merger")
    _theme_file = os.path.join(_theme_dir, "theme.json")

    def _load_saved_theme():
        try:
            with open(_theme_file, "r") as f:
                data = _json.load(f)
            name = data.get("theme", "Dark")
            if name in THEMES:
                return name
        except Exception:
            pass
        return "Dark"

    def _save_theme(name):
        try:
            os.makedirs(_theme_dir, exist_ok=True)
            with open(_theme_file, "w") as f:
                _json.dump({"theme": name}, f)
        except Exception:
            pass

    current_theme_name = _load_saved_theme()

    # ── fonts ──
    FONT     = ("Segoe UI", 10)
    FONT_SM  = ("Segoe UI", 9)
    FONT_LOG = ("Consolas", 9)
    FONT_H   = ("Segoe UI", 13, "bold")
    FONT_SECTION = ("Segoe UI", 10, "bold")
    FONT_MERGE = ("Segoe UI", 11, "bold")

    root = tk.Tk()
    root.title("Simple eBook Merger v3.0")
    root.resizable(True, True)
    root.geometry("760x920")
    root.minsize(720, 700)

    style = ttk.Style()
    style.theme_use("clam")

    # keep refs to tk widgets that need manual recoloring (not ttk)
    tk_widgets = []  # filled in as we build the UI

    def _apply_theme(name):
        """Recolour the entire window to match the named theme."""
        nonlocal current_theme_name
        t = THEMES[name]
        current_theme_name = name

        # ttk styles
        style.configure("TFrame", background=t["bg"])
        style.configure("Mid.TFrame", background=t["bg_mid"])
        style.configure("TLabel", background=t["bg"], foreground=t["fg"], font=FONT)
        style.configure("Dim.TLabel", background=t["bg"], foreground=t["fg_dim"],
                         font=FONT_SM)
        style.configure("Head.TLabel", background=t["bg"],
                         foreground=t["heading_fg"], font=FONT_H)
        style.configure("Section.TLabel", background=t["bg"],
                         foreground=t["heading_fg"], font=FONT_SECTION)

        style.configure("Accent.TButton", font=FONT_MERGE, padding=(16, 8))
        style.map("Accent.TButton",
            background=[("active", t["accent2"]), ("!active", t["accent"])],
            foreground=[("active", "#ffffff"), ("!active", "#ffffff")],
        )
        style.configure("TButton", font=FONT, padding=(10, 5),
                         background=t["button_bg"], foreground=t["button_fg"])
        style.map("TButton",
            background=[("active", t["accent"]), ("!active", t["button_bg"])],
            foreground=[("active", "#ffffff"), ("!active", t["button_fg"])],
        )

        # theme selector buttons
        style.configure("ThemeBtn.TButton", font=FONT_SM, padding=(8, 3),
                         background=t["button_bg"], foreground=t["button_fg"])
        style.map("ThemeBtn.TButton",
            background=[("active", t["accent"]), ("!active", t["button_bg"])],
            foreground=[("active", "#ffffff"), ("!active", t["button_fg"])],
        )
        style.configure("ThemeBtnActive.TButton", font=FONT_SM, padding=(8, 3),
                         background=t["accent"], foreground="#ffffff")
        style.map("ThemeBtnActive.TButton",
            background=[("active", t["accent2"]), ("!active", t["accent"])],
            foreground=[("active", "#ffffff"), ("!active", "#ffffff")],
        )

        style.configure("TCheckbutton", background=t["bg"], foreground=t["fg"],
                         font=FONT)
        style.map("TCheckbutton", background=[("active", t["bg"])])

        style.configure("TRadiobutton", background=t["bg"], foreground=t["fg"],
                         font=FONT)
        style.map("TRadiobutton", background=[("active", t["bg"])])

        style.configure("TEntry", fieldbackground=t["bg_entry"],
                         foreground=t["fg"], insertcolor=t["fg"],
                         font=FONT, padding=5)

        style.configure("TCombobox", fieldbackground=t["bg_entry"],
                         foreground=t["fg"], font=FONT, padding=5)

        style.configure("TLabelframe", background=t["bg"],
                         foreground=t["fg_dim"], bordercolor=t["border"])
        style.configure("TLabelframe.Label", background=t["bg"],
                         foreground=t["fg_dim"], font=FONT_SM)

        style.configure("Status.TLabel", background=t["bg_mid"],
                         foreground=t["fg_dim"], font=FONT_SM)

        # root window
        root.configure(bg=t["bg"])

        # recolour all plain tk widgets we're keeping tabs on
        for w, role in tk_widgets:
            try:
                if role == "listbox":
                    w.configure(bg=t["bg_mid"], fg=t["fg"],
                                selectbackground=t["select_bg"],
                                selectforeground="#fff")
                elif role == "log":
                    w.configure(bg=t["log_bg"], fg=t["log_fg"])
                    w.tag_configure("ok",   foreground=t["green"])
                    w.tag_configure("warn", foreground=t["warn"])
                    w.tag_configure("err",  foreground=t["red"])
                    w.tag_configure("dim",  foreground=t["fg_dim"])
                elif role == "preview":
                    w.configure(bg=t["bg_mid"], fg=t["fg"])
                elif role == "paste_text":
                    w.configure(bg=t["bg_entry"], fg=t["fg"],
                                insertbackground=t["fg"])
                elif role == "statusbar":
                    w.configure(bg=t["bg_mid"])
            except Exception:
                pass

        # update theme-button relief to show which is active
        for tname, btn in theme_buttons.items():
            if tname == name:
                btn.configure(style="ThemeBtnActive.TButton")
            else:
                btn.configure(style="ThemeBtn.TButton")

        _save_theme(name)

    # ── scrollable outer frame ──
    canvas = tk.Canvas(root, highlightthickness=0)
    scrollbar = ttk.Scrollbar(root, orient="vertical", command=canvas.yview)
    canvas.configure(yscrollcommand=scrollbar.set)

    scrollbar.pack(side="right", fill="y")
    canvas.pack(side="left", fill="both", expand=True)

    outer = ttk.Frame(canvas, padding=(16, 10, 16, 0))
    outer_window = canvas.create_window((0, 0), window=outer, anchor="nw")

    def _on_frame_configure(event=None):
        canvas.configure(scrollregion=canvas.bbox("all"))

    def _on_canvas_configure(event):
        canvas.itemconfig(outer_window, width=event.width)

    outer.bind("<Configure>", _on_frame_configure)
    canvas.bind("<Configure>", _on_canvas_configure)

    # mouse wheel scrolling
    def _on_mousewheel(event):
        canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _on_mousewheel_linux(event):
        if event.num == 4:
            canvas.yview_scroll(-1, "units")
        elif event.num == 5:
            canvas.yview_scroll(1, "units")

    canvas.bind_all("<MouseWheel>", _on_mousewheel)
    canvas.bind_all("<Button-4>", _on_mousewheel_linux)
    canvas.bind_all("<Button-5>", _on_mousewheel_linux)

    tk_widgets.append((canvas, "statusbar"))

    # ── theme bar (very top) ──
    theme_bar = ttk.Frame(outer)
    theme_bar.pack(fill="x", pady=(0, 8))
    ttk.Label(theme_bar, text="Theme:", style="Dim.TLabel").pack(
        side="left", padx=(0, 8))

    theme_buttons = {}
    for tname in THEMES:
        btn = ttk.Button(theme_bar, text=tname, style="ThemeBtn.TButton",
                          command=lambda n=tname: _apply_theme(n))
        btn.pack(side="left", padx=(0, 4))
        theme_buttons[tname] = btn

    # dep status — show which optional deps are missing
    _missing_optional = []
    if not HAS_PYPDF:
        _missing_optional.append("pypdf")
    if not HAS_DOCX:
        _missing_optional.append("python-docx")
    if not HAS_ODFPY:
        _missing_optional.append("odfpy")
    if not HAS_FPDF2:
        _missing_optional.append("fpdf2")
    if not HAS_PILLOW:
        _missing_optional.append("Pillow")

    if _missing_optional:
        dep_note = "Optional deps not installed: " + ", ".join(_missing_optional)
        ttk.Label(theme_bar, text=dep_note, style="Dim.TLabel").pack(
            side="right", padx=(8, 0))

    # ── header ──
    ttk.Label(outer, text="Simple eBook Merger v3.0",
              style="Head.TLabel").pack(anchor="w")
    ttk.Label(outer, text="Toss your files in a folder, pick it, hit merge.",
              style="Dim.TLabel").pack(anchor="w", pady=(0, 10))

    # ── profiles section ──
    prof_frame = ttk.LabelFrame(outer, text="  Profiles  ", padding=(10, 6))
    prof_frame.pack(fill="x", pady=(0, 10))

    profile_label_var = tk.StringVar(value="No profile loaded")

    prof_btn_row = ttk.Frame(prof_frame)
    prof_btn_row.pack(fill="x")

    ttk.Label(prof_btn_row, textvariable=profile_label_var,
              style="Dim.TLabel").pack(side="left", padx=(0, 12))

    def _gather_settings():
        """Collect current GUI fields into a dict."""
        return {
            "input_dir": input_var.get().strip(),
            "output_file": output_var.get().strip(),
            "book_title": title_var.get().strip(),
            "book_author": author_var.get().strip(),
            "add_dividers": dividers_var.get(),
            "output_format": format_var.get().lower(),
            "sort_order": _get_sort_mode(),
            "cover_mode": cover_mode_var.get(),
            "cover_path": cover_path_var.get().strip(),
        }

    def _apply_profile(data):
        """Push a profile dict into the GUI fields."""
        input_var.set(data.get("input_dir", ""))
        output_var.set(data.get("output_file", ""))
        title_var.set(data.get("book_title", "My Merged eBook"))
        author_var.set(data.get("book_author", "Various Authors"))
        dividers_var.set(data.get("add_dividers", True))
        fmt = data.get("output_format", "epub").upper()
        if fmt in [f.upper() for f in _available_output_formats()]:
            format_var.set(fmt)
        srt = data.get("sort_order", "smart")
        if srt == "alpha":
            sort_var.set("Alphabetical (A-Z)")
        else:
            sort_var.set("Smart Order (Recommended)")
        cover_mode_var.set(data.get("cover_mode", "none"))
        cover_path_var.set(data.get("cover_path", ""))
        pname = data.get("name", "")
        if pname:
            profile_label_var.set(f"Profile: {pname}")
            root.title(f"Simple eBook Merger v3.0 \u2014 {pname}")
        _refresh_file_list()

    def _save_profile_dialog():
        name = tk.simpledialog.askstring(
            "Save Profile",
            "Give this profile a name:",
            parent=root,
        )
        if not name or not name.strip():
            return
        name = name.strip()
        settings = _gather_settings()
        path = save_profile(name, settings)
        profile_label_var.set(f"Profile: {name}")
        root.title(f"Simple eBook Merger v3.0 \u2014 {name}")
        messagebox.showinfo("Saved", f"Profile \"{name}\" saved.")

    def _load_profile_dialog():
        profs = list_profiles()
        if not profs:
            messagebox.showinfo("No Profiles",
                "No saved profiles yet.\nUse \"Save Profile\" to create one.")
            return
        names = [p[0] for p in profs]
        _pick_profile_window(names)

    def _popup_colors():
        """Grab current theme colours for popup windows."""
        t = THEMES[current_theme_name]
        return t

    def _pick_profile_window(names):
        """Small popup to pick a profile from a list."""
        t = _popup_colors()
        win = tk.Toplevel(root)
        win.title("Load Profile")
        win.configure(bg=t["bg"])
        win.geometry("340x300")
        win.transient(root)
        win.grab_set()

        ttk.Label(win, text="Pick a profile to load:").pack(
            anchor="w", padx=12, pady=(12, 6))

        listbox = tk.Listbox(win, bg=t["bg_mid"], fg=t["fg"], font=FONT_SM,
                              selectbackground=t["accent"],
                              selectforeground="#fff",
                              borderwidth=0, highlightthickness=1,
                              highlightcolor=t["accent"])
        listbox.pack(fill="both", expand=True, padx=12, pady=(0, 8))
        for n in names:
            listbox.insert("end", f"  {n}")

        def _do_load():
            sel = listbox.curselection()
            if not sel:
                return
            chosen = names[sel[0]]
            data = load_profile(chosen)
            if data:
                _apply_profile(data)
            win.destroy()

        btn_row = ttk.Frame(win)
        btn_row.pack(fill="x", padx=12, pady=(0, 12))
        ttk.Button(btn_row, text="Load", command=_do_load).pack(
            side="right", padx=(6, 0))
        ttk.Button(btn_row, text="Cancel", command=win.destroy).pack(
            side="right")

        listbox.bind("<Double-1>", lambda e: _do_load())

    def _delete_profile_dialog():
        profs = list_profiles()
        if not profs:
            messagebox.showinfo("No Profiles",
                "No saved profiles to delete.")
            return
        names = [p[0] for p in profs]

        t = _popup_colors()
        win = tk.Toplevel(root)
        win.title("Delete Profile")
        win.configure(bg=t["bg"])
        win.geometry("340x300")
        win.transient(root)
        win.grab_set()

        ttk.Label(win, text="Pick a profile to delete:").pack(
            anchor="w", padx=12, pady=(12, 6))

        listbox = tk.Listbox(win, bg=t["bg_mid"], fg=t["fg"], font=FONT_SM,
                              selectbackground=t["red"],
                              selectforeground="#fff",
                              borderwidth=0, highlightthickness=1,
                              highlightcolor=t["red"])
        listbox.pack(fill="both", expand=True, padx=12, pady=(0, 8))
        for n in names:
            listbox.insert("end", f"  {n}")

        def _do_delete():
            sel = listbox.curselection()
            if not sel:
                return
            chosen = names[sel[0]]
            if messagebox.askyesno("Confirm",
                    f"Delete profile \"{chosen}\"?", parent=win):
                delete_profile(chosen)
                listbox.delete(sel[0])
                names.pop(sel[0])
                # if that was the loaded profile, clear the label
                if chosen in profile_label_var.get():
                    profile_label_var.set("No profile loaded")
                    root.title("Simple eBook Merger v3.0")
                if not names:
                    win.destroy()

        btn_row = ttk.Frame(win)
        btn_row.pack(fill="x", padx=12, pady=(0, 12))
        ttk.Button(btn_row, text="Delete", command=_do_delete).pack(
            side="right", padx=(6, 0))
        ttk.Button(btn_row, text="Close", command=win.destroy).pack(
            side="right")

    # we need simpledialog for the save-profile name prompt
    import tkinter.simpledialog as _sd
    tk.simpledialog = _sd

    ttk.Button(prof_btn_row, text="Save Profile",
               command=_save_profile_dialog).pack(side="right", padx=(4, 0))
    ttk.Button(prof_btn_row, text="Load Profile",
               command=_load_profile_dialog).pack(side="right", padx=(4, 0))
    ttk.Button(prof_btn_row, text="Delete Profile",
               command=_delete_profile_dialog).pack(side="right", padx=(4, 0))

    # ── input section ──
    input_frame = ttk.LabelFrame(outer, text="  Input  ", padding=(10, 6))
    input_frame.pack(fill="x", pady=(0, 10))

    row1 = ttk.Frame(input_frame)
    row1.pack(fill="x", pady=(0, 6))
    ttk.Label(row1, text="Folder:").pack(side="left")

    input_var = tk.StringVar()
    input_entry = ttk.Entry(row1, textvariable=input_var)
    input_entry.pack(side="left", fill="x", expand=True, padx=(8, 6))

    def pick_folder():
        d = filedialog.askdirectory(title="Pick folder with your eBooks")
        if d:
            input_var.set(d)
            _refresh_file_list()

    ttk.Button(row1, text="Browse...", command=pick_folder).pack(side="right")

    # ── file list + buttons row ──
    file_btn_row = ttk.Frame(input_frame)
    file_btn_row.pack(fill="x", pady=(0, 4))

    # file finder / scan folder button
    def _scan_folder_dialog():
        """Open a folder picker, scan recursively, let user pick files."""
        d = filedialog.askdirectory(title="Pick folder to scan recursively")
        if not d:
            return
        found = find_ebook_files_recursive(d)
        if not found:
            messagebox.showinfo("Nothing Found",
                "No supported ebook files found in that folder tree.")
            return

        t = _popup_colors()
        win = tk.Toplevel(root)
        win.title("Found Files")
        win.configure(bg=t["bg"])
        win.geometry("500x400")
        win.transient(root)
        win.grab_set()

        ttk.Label(win, text=f"Found {len(found)} file(s) — select which to add:").pack(
            anchor="w", padx=12, pady=(12, 6))

        listbox = tk.Listbox(win, bg=t["bg_mid"], fg=t["fg"], font=FONT_SM,
                              selectmode="extended",
                              selectbackground=t["accent"],
                              selectforeground="#fff",
                              borderwidth=0, highlightthickness=1,
                              highlightcolor=t["accent"])
        listbox.pack(fill="both", expand=True, padx=12, pady=(0, 8))
        for full, rel in found:
            listbox.insert("end", f"  {rel}")

        def _select_all():
            listbox.select_set(0, "end")

        def _do_add():
            sel = listbox.curselection()
            if not sel:
                return
            for idx in sel:
                full, rel = found[idx]
                name = os.path.basename(full)
                # add to our tracked files list
                _add_individual_file(full, name)
            win.destroy()

        btn_row = ttk.Frame(win)
        btn_row.pack(fill="x", padx=12, pady=(0, 12))
        ttk.Button(btn_row, text="Add Selected", command=_do_add).pack(
            side="right", padx=(6, 0))
        ttk.Button(btn_row, text="Select All", command=_select_all).pack(
            side="right", padx=(6, 0))
        ttk.Button(btn_row, text="Cancel", command=win.destroy).pack(
            side="right")

    ttk.Button(file_btn_row, text="Scan Folder...",
               command=_scan_folder_dialog).pack(side="left", padx=(0, 4))

    # paste paths button
    paste_visible = tk.BooleanVar(value=False)

    def _toggle_paste():
        if paste_visible.get():
            paste_frame.pack_forget()
            paste_visible.set(False)
        else:
            paste_frame.pack(fill="x", pady=(4, 4), after=file_btn_row)
            paste_visible.set(True)

    ttk.Button(file_btn_row, text="Paste Paths...",
               command=_toggle_paste).pack(side="left", padx=(0, 4))

    # paste paths text area (hidden by default)
    paste_frame = ttk.Frame(input_frame)
    paste_text = tk.Text(paste_frame, font=FONT_SM, height=4,
                          wrap="word", borderwidth=1)
    paste_text.pack(fill="x", side="left", expand=True)
    tk_widgets.append((paste_text, "paste_text"))

    def _do_paste_add():
        """Add files from the paste text box."""
        raw = paste_text.get("1.0", "end").strip()
        if not raw:
            return
        added = 0
        for line in raw.split("\n"):
            line = line.strip().strip('"').strip("'")
            if not line:
                continue
            if os.path.isfile(line):
                ext = os.path.splitext(line)[1].lower()
                if ext in ALLOWED_EXTENSIONS:
                    _add_individual_file(line, os.path.basename(line))
                    added += 1
        paste_text.delete("1.0", "end")
        if added == 0:
            messagebox.showinfo("Nothing Added",
                "No valid ebook file paths found in the pasted text.")

    ttk.Button(paste_frame, text="Add", command=_do_paste_add).pack(
        side="right", padx=(4, 0))

    # ── individual files list (not folder-based) ──
    # we track files from both the folder scanner and paste/scan features
    _individual_files = []  # list of (path, name) tuples

    def _add_individual_file(path, name):
        """Add a file to the individual files list + refresh display."""
        # don't add duplicates
        for p, _ in _individual_files:
            if os.path.abspath(p) == os.path.abspath(path):
                return
        _individual_files.append((path, name))
        _refresh_file_list()

    # file preview inside the input section
    file_listbox = tk.Listbox(input_frame, font=FONT_SM,
                               borderwidth=0, highlightthickness=0, height=8)
    file_listbox.pack(fill="both", expand=True, pady=(0, 4))
    tk_widgets.append((file_listbox, "listbox"))

    file_count_var = tk.StringVar(value="No folder selected yet")
    ttk.Label(input_frame, textvariable=file_count_var,
              style="Dim.TLabel").pack(anchor="w")

    # ── chapter preview panel ──
    preview_frame = ttk.LabelFrame(input_frame, text="  Preview  ", padding=(6, 4))
    preview_frame.pack(fill="x", pady=(4, 0))

    preview_text = tk.Text(preview_frame, font=FONT_SM, wrap="word",
                            borderwidth=0, highlightthickness=0,
                            state="disabled", height=5)
    preview_scroll = ttk.Scrollbar(preview_frame, orient="vertical",
                                    command=preview_text.yview)
    preview_text.configure(yscrollcommand=preview_scroll.set)
    preview_scroll.pack(side="right", fill="y")
    preview_text.pack(fill="both", expand=True)
    tk_widgets.append((preview_text, "preview"))

    # ── track current file list for the merge + preview ──
    _current_files = []  # the files that will actually be merged

    def _on_file_select(event=None):
        """Update chapter preview when user clicks a file in the list."""
        sel = file_listbox.curselection()
        if not sel or not _current_files:
            return
        idx = sel[0]
        if idx >= len(_current_files):
            return
        fpath, fname = _current_files[idx]

        # run preview in a thread so we don't freeze the GUI
        def _do_preview():
            text = _preview_file_text(fpath)
            def _update():
                preview_text.configure(state="normal")
                preview_text.delete("1.0", "end")
                preview_text.insert("1.0", text)
                preview_text.configure(state="disabled")
            root.after(0, _update)

        threading.Thread(target=_do_preview, daemon=True).start()

    file_listbox.bind("<<ListboxSelect>>", _on_file_select)

    # ── output section ──
    output_frame = ttk.LabelFrame(outer, text="  Output  ", padding=(10, 6))
    output_frame.pack(fill="x", pady=(0, 10))

    row2 = ttk.Frame(output_frame)
    row2.pack(fill="x", pady=(0, 6))
    ttk.Label(row2, text="Save as:").pack(side="left")

    output_var = tk.StringVar(value="merged_book.epub")
    output_entry = ttk.Entry(row2, textvariable=output_var)
    output_entry.pack(side="left", fill="x", expand=True, padx=(8, 6))

    def pick_output():
        fmt = format_var.get().lower()
        ext = OUTPUT_EXT_MAP.get(fmt, ".epub")
        ftypes = [
            ("EPUB files", "*.epub"),
            ("Text files", "*.txt"),
            ("HTML files", "*.html"),
            ("PDF files", "*.pdf"),
            ("Word documents", "*.docx"),
            ("All files", "*.*"),
        ]
        f = filedialog.asksaveasfilename(
            title="Save merged file as",
            defaultextension=ext,
            filetypes=ftypes,
        )
        if f:
            output_var.set(f)

    ttk.Button(row2, text="Browse...", command=pick_output).pack(side="right")

    # format + title + author in a sub-grid
    meta_row = ttk.Frame(output_frame)
    meta_row.pack(fill="x", pady=(0, 2))

    ttk.Label(meta_row, text="Format:").pack(side="left")
    avail_fmts = [f.upper() for f in _available_output_formats()]
    format_var = tk.StringVar(value="EPUB")
    format_combo = ttk.Combobox(
        meta_row, textvariable=format_var,
        values=avail_fmts,
        state="readonly", width=8,
    )
    format_combo.pack(side="left", padx=(8, 16))

    def _on_format_change(event=None):
        fmt = format_var.get().lower()
        ext = OUTPUT_EXT_MAP.get(fmt, ".epub")
        current = output_var.get().strip()
        if current:
            base = os.path.splitext(current)[0]
            output_var.set(base + ext)

    format_combo.bind("<<ComboboxSelected>>", _on_format_change)

    ttk.Label(meta_row, text="Title:").pack(side="left")
    title_var = tk.StringVar(value="My Merged eBook")
    ttk.Entry(meta_row, textvariable=title_var, width=20).pack(
        side="left", padx=(8, 16))

    ttk.Label(meta_row, text="Author:").pack(side="left")
    author_var = tk.StringVar(value="Various Authors")
    ttk.Entry(meta_row, textvariable=author_var).pack(
        side="left", fill="x", expand=True, padx=(8, 0))

    # ── cover image section ──
    cover_frame = ttk.LabelFrame(outer, text="  Cover Image  ", padding=(10, 6))
    cover_frame.pack(fill="x", pady=(0, 10))

    cover_mode_var = tk.StringVar(value="none")
    cover_path_var = tk.StringVar(value="")

    cover_opt_row = ttk.Frame(cover_frame)
    cover_opt_row.pack(fill="x")

    ttk.Radiobutton(cover_opt_row, text="None", variable=cover_mode_var,
                     value="none").pack(side="left", padx=(0, 12))
    ttk.Radiobutton(cover_opt_row, text="Auto (from first input)",
                     variable=cover_mode_var, value="auto").pack(
                         side="left", padx=(0, 12))
    ttk.Radiobutton(cover_opt_row, text="Browse...", variable=cover_mode_var,
                     value="browse").pack(side="left", padx=(0, 8))

    cover_path_entry = ttk.Entry(cover_opt_row, textvariable=cover_path_var,
                                  width=30)
    cover_path_entry.pack(side="left", fill="x", expand=True, padx=(0, 6))

    def _pick_cover():
        f = filedialog.askopenfilename(
            title="Pick a cover image",
            filetypes=[
                ("Image files", "*.png *.jpg *.jpeg *.bmp *.gif *.webp"),
                ("All files", "*.*"),
            ],
        )
        if f:
            cover_path_var.set(f)
            cover_mode_var.set("browse")

    ttk.Button(cover_opt_row, text="Pick...", command=_pick_cover).pack(
        side="right")

    if not HAS_PILLOW:
        ttk.Label(cover_frame, text="(Cover for PDF output needs Pillow)",
                  style="Dim.TLabel").pack(anchor="w", pady=(4, 0))

    # ── options section ──
    options_frame = ttk.LabelFrame(outer, text="  Options  ", padding=(10, 6))
    options_frame.pack(fill="x", pady=(0, 10))

    opt_row = ttk.Frame(options_frame)
    opt_row.pack(fill="x")

    dividers_var = tk.BooleanVar(value=True)
    ttk.Checkbutton(opt_row, text="Insert divider pages between files",
                     variable=dividers_var).pack(side="left")

    ttk.Label(opt_row, text="Sort:").pack(side="left", padx=(24, 0))
    sort_var = tk.StringVar(value="Smart Order (Recommended)")
    sort_combo = ttk.Combobox(
        opt_row, textvariable=sort_var,
        values=["Smart Order (Recommended)", "Alphabetical (A-Z)"],
        state="readonly", width=26,
    )
    sort_combo.pack(side="left", padx=(8, 0))

    def _on_sort_change(event=None):
        _refresh_file_list()

    sort_combo.bind("<<ComboboxSelected>>", _on_sort_change)

    def _get_sort_mode():
        return "smart" if "Smart" in sort_var.get() else "alpha"

    def _refresh_file_list():
        nonlocal _current_files
        file_listbox.delete(0, "end")
        folder = input_var.get().strip()

        # gather files from folder + individual adds
        combined = []
        if folder and os.path.isdir(folder):
            mode = _get_sort_mode()
            found, skipped, sort_summary = find_ebook_files(folder, sort_order=mode)
            combined.extend(found)

        # add individual files that aren't already in the list
        folder_paths = {os.path.abspath(p) for p, _ in combined}
        for p, n in _individual_files:
            if os.path.abspath(p) not in folder_paths:
                combined.append((p, n))

        _current_files = combined

        for _, name in combined:
            ext = os.path.splitext(name)[1].lower()
            file_listbox.insert("end", f"  {name}   ({ext})")

        note = f"{len(combined)} file(s)"
        if not combined and not folder:
            note = "No folder selected yet"
        file_count_var.set(note)

        # clear preview
        preview_text.configure(state="normal")
        preview_text.delete("1.0", "end")
        preview_text.configure(state="disabled")

    # ── merge section ──
    merge_frame = ttk.LabelFrame(outer, text="  Merge  ", padding=(10, 8))
    merge_frame.pack(fill="x", pady=(0, 10))

    merge_row = ttk.Frame(merge_frame)
    merge_row.pack(fill="x")

    status_var = tk.StringVar(value="Ready")
    ttk.Label(merge_row, textvariable=status_var,
              style="Dim.TLabel").pack(side="left")

    merge_btn = ttk.Button(merge_row, text="  Merge  ",
                            style="Accent.TButton")
    merge_btn.pack(side="right")

    # ── history section ──
    history_frame = ttk.LabelFrame(outer, text="  History  ", padding=(10, 6))
    history_frame.pack(fill="x", pady=(0, 10))

    history_listbox = tk.Listbox(history_frame, font=FONT_SM,
                                  borderwidth=0, highlightthickness=0, height=5)
    history_listbox.pack(fill="both", expand=True, pady=(0, 4))
    tk_widgets.append((history_listbox, "listbox"))

    history_btn_row = ttk.Frame(history_frame)
    history_btn_row.pack(fill="x")

    _history_entries = []

    def _refresh_history():
        nonlocal _history_entries
        history_listbox.delete(0, "end")
        _history_entries = _load_history()
        # show last 20, newest first
        for entry in reversed(_history_entries[-20:]):
            ts = entry.get("timestamp", "?")
            fmt = entry.get("output_format", "?").upper()
            n_files = len(entry.get("input_files", []))
            title = entry.get("title", "?")
            history_listbox.insert("end",
                f"  {ts}  |  {title}  |  {fmt}  |  {n_files} file(s)")

    def _rerun_history():
        """Reload settings from a selected history entry."""
        sel = history_listbox.curselection()
        if not sel:
            messagebox.showinfo("Select Entry",
                "Click on a history entry first.")
            return
        # entries are displayed newest-first, but stored oldest-first
        display_idx = sel[0]
        entries = _history_entries[-20:]
        actual_idx = len(entries) - 1 - display_idx
        if actual_idx < 0 or actual_idx >= len(entries):
            return
        entry = entries[actual_idx]

        # populate fields from the history entry
        title_var.set(entry.get("title", "My Merged eBook"))
        author_var.set(entry.get("author", "Various Authors"))
        dividers_var.set(entry.get("dividers", True))
        fmt = entry.get("output_format", "epub").upper()
        if fmt in avail_fmts:
            format_var.set(fmt)
        output_var.set(entry.get("output_path", "merged_book.epub"))
        srt = entry.get("sort_order", "smart")
        if srt == "alpha":
            sort_var.set("Alphabetical (A-Z)")
        else:
            sort_var.set("Smart Order (Recommended)")
        cover_mode_var.set(entry.get("cover_mode", "none"))
        cover_path_var.set(entry.get("cover_path", ""))

        # try to reload the input files
        _individual_files.clear()
        input_paths = entry.get("input_paths", [])
        input_names = entry.get("input_files", [])
        for i, p in enumerate(input_paths):
            if os.path.isfile(p):
                n = input_names[i] if i < len(input_names) else os.path.basename(p)
                _individual_files.append((p, n))

        # clear the folder input since we're using individual files
        input_var.set("")
        _refresh_file_list()

    def _undo_last():
        """Delete the output file from the most recent merge."""
        if not _history_entries:
            messagebox.showinfo("No History", "No merge history yet.")
            return
        last = _history_entries[-1]
        out_path = last.get("output_path", "")
        if not out_path or not os.path.isfile(out_path):
            messagebox.showinfo("File Not Found",
                f"Output file doesn't exist:\n{out_path}")
            return
        if messagebox.askyesno("Undo Last Merge",
                f"Delete this output file?\n\n{out_path}"):
            try:
                os.remove(out_path)
                messagebox.showinfo("Done", "File deleted.")
            except OSError as e:
                messagebox.showerror("Error", f"Couldn't delete: {e}")

    ttk.Button(history_btn_row, text="Re-run Selected",
               command=_rerun_history).pack(side="left", padx=(0, 6))
    ttk.Button(history_btn_row, text="Undo Last (Delete File)",
               command=_undo_last).pack(side="left", padx=(0, 6))
    ttk.Button(history_btn_row, text="Refresh",
               command=_refresh_history).pack(side="right")

    # load history on startup
    _refresh_history()

    # ── log section ──
    log_frame = ttk.LabelFrame(outer, text="  Log  ", padding=(8, 6))
    log_frame.pack(fill="both", expand=True, pady=(0, 8))

    log_text = tk.Text(log_frame, font=FONT_LOG,
                        wrap="word", borderwidth=0, highlightthickness=0,
                        state="disabled", height=10)
    log_scroll = ttk.Scrollbar(log_frame, orient="vertical",
                                command=log_text.yview)
    log_text.configure(yscrollcommand=log_scroll.set)
    log_scroll.pack(side="right", fill="y")
    log_text.pack(side="left", fill="both", expand=True)
    tk_widgets.append((log_text, "log"))

    def write_log(msg, tag=None):
        log_text.configure(state="normal")
        if tag:
            log_text.insert("end", msg + "\n", tag)
        else:
            log_text.insert("end", msg + "\n")
        log_text.see("end")
        log_text.configure(state="disabled")

    # ── status bar at the very bottom ──
    status_bar = tk.Frame(root, height=24)
    status_bar.pack(fill="x", side="bottom")
    tk_widgets.append((status_bar, "statusbar"))

    status_bar_label = ttk.Label(status_bar, textvariable=status_var,
                                  style="Status.TLabel", padding=(8, 3))
    status_bar_label.pack(fill="x")

    # ── apply the saved theme now that all widgets exist ──
    _apply_theme(current_theme_name)

    # ── merge action ──
    def do_merge():
        out    = output_var.get().strip()
        ttl    = title_var.get().strip() or "Merged eBook"
        auth   = author_var.get().strip() or "Unknown"
        divs   = dividers_var.get()
        fmt    = format_var.get().lower()
        cmode  = cover_mode_var.get()
        cpath  = cover_path_var.get().strip()

        # clear old log
        log_text.configure(state="normal")
        log_text.delete("1.0", "end")
        log_text.configure(state="disabled")

        if not out:
            messagebox.showerror("Oops", "Set an output file path.")
            return

        files = _current_files
        if not files:
            folder = input_var.get().strip()
            if folder and os.path.isdir(folder):
                mode = _get_sort_mode()
                files, skipped, sort_summary = find_ebook_files(folder, sort_order=mode)
            if not files:
                exts = ", ".join(sorted(ALLOWED_EXTENSIONS))
                messagebox.showwarning("Nothing to merge",
                    f"No files to merge. Add files or pick a folder first.")
                return

        # disable the button so you can't double-click
        merge_btn.configure(state="disabled")
        status_var.set("Merging...")

        mode = _get_sort_mode()
        sort_label = "Smart order" if mode == "smart" else "Alphabetical"
        write_log(f"  Sort: {sort_label}", "dim")
        write_log("")

        def _run():
            try:
                stats = merge_files(
                    files, out, ttl, auth,
                    dividers=divs,
                    output_format=fmt,
                    cover_mode=cmode,
                    cover_path=cpath,
                    log=lambda m: root.after(0, write_log, m),
                )

                # save to history
                _add_history_entry(files, out, fmt, ttl, auth, divs,
                                   mode, cmode, cpath)

                def _done():
                    write_log("")
                    write_log(f"  Done! {stats['files_merged']} file(s), "
                              f"{stats['chapters']} chapter(s), "
                              f"~{stats['words']:,} words, {stats['size']}", "ok")
                    write_log(f"  Saved \u2192 {stats['path']}", "ok")
                    status_var.set("Done!")
                    merge_btn.configure(state="normal")
                    _refresh_history()
                root.after(0, _done)

            except Exception as exc:
                def _fail():
                    write_log(f"\n  ERROR: {exc}", "err")
                    status_var.set("Failed")
                    merge_btn.configure(state="normal")
                root.after(0, _fail)

        # run in a thread so the GUI doesn't freeze
        threading.Thread(target=_run, daemon=True).start()

    merge_btn.configure(command=do_merge)

    # ── go ──
    root.mainloop()


# ──────────────────────────────────────────────────────────────
# Terminal (headless) mode — works the same as before, plus
# output format support
# ──────────────────────────────────────────────────────────────

def run_terminal():
    print("\n" + "=" * 60)
    print("  Simple eBook Merger v3.0")
    print("=" * 60 + "\n")

    print("\u26a0  LEGAL REMINDER:")
    print("   Personal use only. Only merge files you actually own.")
    print("   Don't distribute copyrighted material.")
    print("   I am not liable for ANYTHING. Not your files, not your")
    print("   computer, not your cat, not your existential dread.")
    print("   By running this you accept full responsibility and agree")
    print("   to the full disclaimer in the source code. Read it.\n")

    # show dep status
    _opt_deps = []
    if not HAS_PYPDF:
        _opt_deps.append("pypdf (PDF input)")
    if not HAS_DOCX:
        _opt_deps.append("python-docx (DOCX)")
    if not HAS_ODFPY:
        _opt_deps.append("odfpy (ODT)")
    if not HAS_FPDF2:
        _opt_deps.append("fpdf2 (PDF output)")
    if not HAS_PILLOW:
        _opt_deps.append("Pillow (CBZ + covers)")
    if _opt_deps:
        print("  Optional deps not installed:")
        for d in _opt_deps:
            print(f"    - {d}")
        print()

    # ── handle --list-profiles ──
    if "--list-profiles" in sys.argv:
        profs = list_profiles()
        if not profs:
            print("  No saved profiles yet.\n")
            print("  Save one with:  --save-profile \"My Profile Name\"\n")
        else:
            print(f"  Saved profiles ({len(profs)}):\n")
            for pname, ppath in profs:
                try:
                    with open(ppath, "r", encoding="utf-8") as fh:
                        pdata = json.load(fh)
                    fmt = pdata.get("output_format", "?").upper()
                    inp = pdata.get("input_dir", "?")
                    updated = pdata.get("updated", "?")
                    print(f"    \u2022 {pname}")
                    print(f"      format: {fmt}  |  input: {inp}")
                    print(f"      last updated: {updated}")
                    print()
                except (json.JSONDecodeError, OSError):
                    print(f"    \u2022 {pname}  (couldn't read file)")
                    print()
        return

    # ── handle --save-profile ──
    if "--save-profile" in sys.argv:
        idx = sys.argv.index("--save-profile")
        if idx + 1 >= len(sys.argv):
            print("  \u2717 --save-profile needs a name.\n")
            print("  Example:  --save-profile \"My Fantasy Series\"\n")
            return
        prof_name = sys.argv[idx + 1]
        settings = {
            "input_dir": INPUT_DIR,
            "output_file": OUTPUT_FILE,
            "book_title": BOOK_TITLE,
            "book_author": BOOK_AUTHOR,
            "add_dividers": ADD_DIVIDERS,
            "output_format": OUTPUT_FORMAT,
            "sort_order": SORT_ORDER,
        }
        path = save_profile(prof_name, settings)
        print(f"  \u2713 Profile saved: \"{prof_name}\"")
        print(f"    \u2192 {path}\n")
        return

    # ── load profile if --profile was given ──
    input_dir = INPUT_DIR
    output_file = OUTPUT_FILE
    book_title = BOOK_TITLE
    book_author = BOOK_AUTHOR
    add_dividers = ADD_DIVIDERS
    output_format = OUTPUT_FORMAT
    sort_order = SORT_ORDER

    if "--profile" in sys.argv:
        idx = sys.argv.index("--profile")
        if idx + 1 >= len(sys.argv):
            print("  \u2717 --profile needs a name.\n")
            print("  Example:  --profile \"My Fantasy Series\"")
            print("  See saved profiles with:  --list-profiles\n")
            return
        prof_name = sys.argv[idx + 1]
        pdata = load_profile(prof_name)
        if pdata is None:
            print(f"  \u2717 No profile found named \"{prof_name}\".\n")
            print("  See saved profiles with:  --list-profiles\n")
            return
        input_dir = pdata.get("input_dir", input_dir)
        output_file = pdata.get("output_file", output_file)
        book_title = pdata.get("book_title", book_title)
        book_author = pdata.get("book_author", book_author)
        add_dividers = pdata.get("add_dividers", add_dividers)
        output_format = pdata.get("output_format", output_format)
        sort_order = pdata.get("sort_order", sort_order)
        print(f"  Loaded profile: \"{prof_name}\"\n")

    # check config
    problems = []
    if "PUT_YOUR_FOLDER" in input_dir or not input_dir.strip():
        problems.append("INPUT_DIR isn't set \u2014 edit the CONFIG section at the top of the script.")
    elif not os.path.isdir(input_dir):
        problems.append(f"INPUT_DIR doesn't exist: {input_dir}")
    if not output_file.strip():
        problems.append("OUTPUT_FILE isn't set.")

    fmt = output_format.lower().strip()
    if fmt not in _available_output_formats():
        avail = ", ".join(_available_output_formats())
        problems.append(f"OUTPUT_FORMAT '{output_format}' isn't available. "
                        f"Available: {avail}")

    sort_mode = sort_order.lower().strip() if isinstance(sort_order, str) else "smart"
    if sort_mode not in ("smart", "alpha"):
        problems.append(f"SORT_ORDER '{sort_order}' isn't valid. Use 'smart' or 'alpha'.")
        sort_mode = "smart"

    if problems:
        print("=" * 60)
        print("  Setup incomplete:")
        print("=" * 60)
        for p in problems:
            print(f"\n  \u2717 {p}")
        print("\nOpen this script in a text editor and fill in the CONFIG section.\n")
        return

    print(f"\u2192 Scanning: {input_dir}\n")
    files, skipped, sort_summary = find_ebook_files(input_dir, sort_order=sort_mode)

    for s in skipped:
        print(f"  \u26a0 skipping: {s}")

    if not files:
        exts = ", ".join(sorted(_available_input_extensions()))
        print(f"\nNo supported files found.")
        print(f"  Available formats: {exts}")
        print("  Double-check INPUT_DIR.\n")
        return

    sort_label = "Smart order" if sort_mode == "smart" else "Alphabetical"
    print(f"  Found {len(files)} file(s) \u2014 {sort_label} ({sort_summary}):")
    for _, name in files:
        print(f"    \u2022 {name}")

    print(f"\n\u2192 Merging into {fmt.upper()} format...\n")

    try:
        stats = merge_files(files, output_file, book_title, book_author,
                            dividers=add_dividers, output_format=fmt)

        # save to history
        _add_history_entry(files, output_file, fmt, book_title, book_author,
                           add_dividers, sort_mode)

        print(f"\n" + "=" * 60)
        print(f"  Files merged   : {stats['files_merged']}")
        print(f"  Chapters total : {stats['chapters']}")
        print(f"  Total words    : ~{stats['words']:,}")
        print(f"  Output size    : {stats['size']}")
        print(f"  Output format  : {fmt.upper()}")
        print("=" * 60)
        print(f"\n\u2713 Done! Saved to:\n  {stats['path']}\n")

    except PermissionError:
        print(f"\nERROR: Can't write to {output_file}")
        print("  Maybe it's open in another program?\n")
    except RuntimeError as e:
        print(f"\n  \u2717 {e}\n")
    except Exception as e:
        print(f"\nSomething went wrong: {e}")
        raise


# ──────────────────────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    # profile management commands work with --nogui or standalone
    if any(a in sys.argv for a in ("--list-profiles", "--save-profile", "--profile")):
        run_terminal()
    elif "--nogui" in sys.argv:
        run_terminal()
    else:
        try:
            launch_gui()
        except ImportError:
            # tkinter not installed (some minimal linux setups)
            print("\n  tkinter isn't available \u2014 falling back to terminal mode.")
            print("  (install it with: sudo apt install python3-tk)\n")
            run_terminal()
